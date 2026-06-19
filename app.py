import os
import argparse
import logging
import json
import sqlite3
import re
import io
import platform
import shutil
import socket
import ipaddress
import glob
import uuid
import hashlib
import secrets
import random
import smtplib
import ssl as _ssl
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timezone, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, g, has_request_context
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect, CSRFError
from flask_talisman import Talisman
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
import ollama
from groq import Groq
import requests
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
from lxml import html as lxml_html

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ADMIN_EMAIL = 'test@marzec.eu'
ADMIN_EMAIL = 'wmarzec@gmail.com'
HISTORY_PAGE_SIZE = 50

_YOUTUBE_RE = re.compile(
    r'^https?://(www\.|m\.)?(youtube\.com/(watch|shorts|embed|live)|youtu\.be/)',
    re.IGNORECASE,
)

_PRIVATE_NETWORKS = [
    ipaddress.ip_network(cidr) for cidr in (
        "127.0.0.0/8", "10.0.0.0/8", "172.16.0.0/12", "192.168.0.0/16",
        "169.254.0.0/16", "0.0.0.0/8", "100.64.0.0/10",
        "::1/128", "fc00::/7", "fe80::/10",
    )
]

def parse_startup_args():
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument(
        '--no-local-models',
        '--skip-local-models',
        '--disable-local-models',
        action='store_true',
        dest='no_local_models'
    )
    args, _ = parser.parse_known_args()
    return args

STARTUP_ARGS = parse_startup_args()
_no_local_models = STARTUP_ARGS.no_local_models or os.getenv('NO_LOCAL_MODELS', '0') == '1'

if _no_local_models:
    whisper = None
else:
    import whisper

def load_env_file(path='.env'):
    if not os.path.exists(path):
        return

    with open(path, 'r', encoding='utf-8') as env_file:
        for raw_line in env_file:
            line = raw_line.strip()
            if not line or line.startswith('#') or '=' not in line:
                continue

            if line.startswith('export '):
                line = line[len('export '):].strip()

            key, value = line.split('=', 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")

            if key and key not in os.environ:
                os.environ[key] = value

load_env_file()

_ENV_FILE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')

def write_env_var(key, value):
    """Update or insert key=value in .env and apply immediately to os.environ."""
    lines = []
    found = False
    if os.path.exists(_ENV_FILE_PATH):
        with open(_ENV_FILE_PATH, 'r', encoding='utf-8') as f:
            for line in f:
                bare = line.strip().lstrip('export').strip()
                if bare.startswith(f'{key}='):
                    lines.append(f'{key}={value}\n')
                    found = True
                else:
                    lines.append(line)
    if not found:
        if lines and not lines[-1].endswith('\n'):
            lines.append('\n')
        lines.append(f'{key}={value}\n')
    with open(_ENV_FILE_PATH, 'w', encoding='utf-8') as f:
        f.writelines(lines)
    if value:
        os.environ[key] = value
    else:
        os.environ.pop(key, None)

from server_restart import restart_server
from models_config import (
    PROVIDERS, MODEL_TYPES, MODEL_CATALOG, LOCAL_MODEL_REQUIREMENTS, DEFAULT_AI_MODELS,
    MODEL_PRICING, AUDIO_PRICING, resolve_pricing_key,
    get_effective_model_pricing, load_pricing_data, save_pricing_data, PRICING_DATA_FILE,
    get_audio_duration_seconds, get_audio_cost,
    get_provider_label, get_provider_api_key, get_provider_env_key, require_provider_api_key,
    raise_invalid_api_key_error, raise_provider_api_error,
    describe_cloud_model, describe_local_notes_model, describe_notes_model, describe_chat_answer_model,
)

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)
app.logger.setLevel(logging.INFO)
_secret_key = os.getenv('FLASK_SECRET_KEY')
if not _secret_key:
    raise RuntimeError(
        "FLASK_SECRET_KEY nie jest ustawiony. "
        "Wygeneruj klucz: python -c \"import secrets; print(secrets.token_hex(32))\" "
        "i dodaj go do pliku .env lub zmiennych środowiskowych."
    )
app.secret_key = _secret_key
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = os.getenv('FLASK_ENV', 'development') == 'production'

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=[],
    storage_uri="memory://",
)
csrf = CSRFProtect(app)

_ytt_api = YouTubeTranscriptApi()

@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    return jsonify({"error": "Nieprawidłowy token CSRF. Odśwież stronę i spróbuj ponownie."}), 400

@app.errorhandler(413)
def handle_too_large(e):
    return jsonify({"error": f"Plik jest zbyt duży. Maksymalny rozmiar to {_max_mb} MB."}), 413

_is_production = os.getenv('FLASK_ENV', 'development') == 'production'
Talisman(
    app,
    force_https=_is_production,
    strict_transport_security=_is_production,
    session_cookie_secure=_is_production,
    content_security_policy={
        'default-src': "'self'",
        'script-src':  "'self' 'unsafe-inline' https://cdn.jsdelivr.net",
        'style-src':   "'self' 'unsafe-inline' https://cdn.jsdelivr.net",
        'font-src':    "'self' https://cdn.jsdelivr.net",
        'img-src':     "'self' data: blob:",
        'media-src':   "'self' blob:",
        'connect-src': "'self' https://cdn.jsdelivr.net",
        'frame-ancestors': "'none'",
    },
    referrer_policy='strict-origin-when-cross-origin',
    feature_policy=False,
    permissions_policy={
        'microphone': '(self)',
        'camera': '()',
        'geolocation': '()',
    },
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'temp_uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

DB_FILE = 'users.db'
LOCAL_MODELS_ENABLED = not _no_local_models
OPENAI_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4o-mini').strip() or 'gpt-4o-mini'
OPENAI_WEB_SEARCH_MODEL = os.getenv('OPENAI_WEB_SEARCH_MODEL', OPENAI_MODEL).strip() or OPENAI_MODEL
OPENAI_TRANSCRIBE_MODEL = os.getenv('OPENAI_TRANSCRIBE_MODEL', 'gpt-4o-mini-transcribe').strip() or 'gpt-4o-mini-transcribe'
OPENAI_USAGE_HISTORY_FILE = os.path.join(BASE_DIR, 'openai_usage_history.jsonl')

_max_mb = int(os.getenv('MAX_UPLOAD_MB', '200'))
app.config['MAX_CONTENT_LENGTH'] = _max_mb * 1024 * 1024

ALLOWED_AUDIO_EXTENSIONS = {
    'wav', 'mp3', 'mp4', 'mpeg', 'mpga', 'm4a', 'webm', 'ogg', 'flac', 'aac', 'opus', 'txt'
}

def validate_password_strength(password):
    """Returns error message string or None if password is strong enough."""
    if not password or len(password) < 8:
        return 'Hasło musi mieć co najmniej 8 znaków.'
    if not any(c.isupper() for c in password):
        return 'Hasło musi zawierać co najmniej jedną wielką literę.'
    if not any(c.isdigit() for c in password):
        return 'Hasło musi zawierać co najmniej jedną cyfrę.'
    if not any(c in '!@#$%^&*(),.?":{}|<>' for c in password):
        return 'Hasło musi zawierać co najmniej jeden znak specjalny (!@#$%^&*...).'
    return None

def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def get_authenticated_user():
    """Zwraca e-mail użytkownika z sesji lub klucza API. None jeśli brak autoryzacji."""
    if 'user_email' in session:
        return session['user_email']
    api_key = None
    auth_header = request.headers.get('Authorization', '')
    if auth_header.startswith('Bearer '):
        api_key = auth_header[7:].strip()
    if not api_key:
        api_key = request.headers.get('X-API-Key', '').strip()
    if not api_key:
        return None
    key_hash = hashlib.sha256(api_key.encode()).hexdigest()
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT user_email, allowed_origin FROM api_keys WHERE key_hash = ?", (key_hash,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return None
    user_email, allowed_origin = row[0], (row[1] or '').strip()
    if allowed_origin:
        base = allowed_origin.rstrip('/')
        origin = request.headers.get('Origin', '').rstrip('/')
        referer = request.headers.get('Referer', '').rstrip('/')
        origin_ok = origin and (origin == base or origin.startswith(base + '/'))
        referer_ok = referer and (referer == base or referer.startswith(base + '/'))
        if not origin_ok and not referer_ok:
            conn.close()
            return None
    cursor.execute("UPDATE api_keys SET last_used_at = CURRENT_TIMESTAMP WHERE key_hash = ?", (key_hash,))
    conn.commit()
    conn.close()
    return user_email

def to_plain_data(value):
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, dict):
        return {str(key): to_plain_data(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [to_plain_data(item) for item in value]
    if hasattr(value, "model_dump"):
        return to_plain_data(value.model_dump())
    if hasattr(value, "to_dict"):
        return to_plain_data(value.to_dict())
    if hasattr(value, "__dict__"):
        data = {
            key: to_plain_data(item)
            for key, item in vars(value).items()
            if not key.startswith("_")
        }
        if data:
            return data

    public_attrs = {}
    for attr_name in dir(value):
        if attr_name.startswith("_"):
            continue
        try:
            attr_value = getattr(value, attr_name)
        except Exception:
            continue
        if callable(attr_value):
            continue
        if isinstance(attr_value, (str, int, float, bool, dict, list, tuple)) or attr_value is None:
            public_attrs[attr_name] = to_plain_data(attr_value)
    if public_attrs:
        return public_attrs

    return str(value)

def get_openai_response_payload(response):
    if response is None:
        return {}
    if isinstance(response, dict):
        return to_plain_data(response)
    if hasattr(response, "json") and callable(response.json):
        try:
            return to_plain_data(response.json())
        except ValueError:
            return {}
    return to_plain_data(response)

def read_field(source, *path):
    current = source
    for key in path:
        if current is None:
            return None
        if isinstance(current, dict):
            current = current.get(key)
        else:
            current = getattr(current, key, None)
    return current

def first_available(source, paths):
    for path in paths:
        value = read_field(source, *path)
        if value is not None:
            return value
    return None

def first_available_from_sources(sources, paths):
    for source in sources:
        value = first_available(source, paths)
        if value is not None:
            return value
    return None

def extract_openai_usage(response, operation_name):
    payload = get_openai_response_payload(response)
    sources = [payload, response]
    usage = first_available_from_sources(sources, [("usage",)])
    usage_data = to_plain_data(usage) if usage is not None else None

    input_tokens = first_available_from_sources(sources, [
        ("usage", "input_tokens"),
        ("usage", "prompt_tokens")
    ])
    output_tokens = first_available_from_sources(sources, [
        ("usage", "output_tokens"),
        ("usage", "completion_tokens")
    ])
    total_tokens = first_available_from_sources(sources, [
        ("usage", "total_tokens")
    ])
    cached_tokens = first_available_from_sources(sources, [
        ("usage", "cached_tokens"),
        ("usage", "input_tokens_details", "cached_tokens"),
        ("usage", "prompt_tokens_details", "cached_tokens")
    ])
    reasoning_tokens = first_available_from_sources(sources, [
        ("usage", "reasoning_tokens"),
        ("usage", "output_tokens_details", "reasoning_tokens"),
        ("usage", "completion_tokens_details", "reasoning_tokens")
    ])
    usage_type = first_available_from_sources(sources, [("usage", "type")])
    seconds = first_available_from_sources(sources, [
        ("usage", "seconds"),
        ("duration",),
    ])

    return {
        "datetime": datetime.now(timezone.utc).isoformat(),
        "operation": operation_name,
        "model": first_available_from_sources(sources, [("model",)]),
        "response_id": first_available_from_sources(sources, [("id",)]),
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": total_tokens,
        "cached_tokens": cached_tokens,
        "reasoning_tokens": reasoning_tokens,
        "usage_type": usage_type,
        "seconds": seconds,
        "usage": usage_data
    }

def get_current_openai_usage_history():
    if not has_request_context():
        return []
    return list(getattr(g, "openai_usage_history", []))

def append_openai_usage_to_request(record):
    if not has_request_context():
        return
    if not hasattr(g, "openai_usage_history"):
        g.openai_usage_history = []
    g.openai_usage_history.append(record)

def append_openai_usage_to_file(record):
    with open(OPENAI_USAGE_HISTORY_FILE, "a", encoding="utf-8") as usage_file:
        usage_file.write(json.dumps(record, ensure_ascii=False) + "\n")

def record_openai_usage(response, operation_name):
    record = extract_openai_usage(response, operation_name)
    append_openai_usage_to_request(record)
    try:
        append_openai_usage_to_file(record)
    except Exception:
        app.logger.exception("Nie udało się zapisać historii użycia OpenAI")
    return record

def serialize_openai_usage_history(records):
    if not records:
        return ""
    return json.dumps(records, ensure_ascii=False)

def parse_openai_usage_history(value):
    if not value:
        return []
    try:
        parsed = json.loads(value)
    except (TypeError, ValueError):
        return []
    return parsed if isinstance(parsed, list) else []

def compute_token_summary(usage_record_lists):
    """Sum input/output tokens across multiple usage-record lists and estimate cost."""
    total_input = 0
    total_output = 0
    cost_usd = 0.0
    for records in usage_record_lists:
        for rec in (records or []):
            inp  = int(rec.get("input_tokens")  or 0)
            out  = int(rec.get("output_tokens") or 0)
            secs = float(rec.get("seconds") or 0)
            total_input  += inp
            total_output += out
            if inp or out:
                pricing_key = resolve_pricing_key(rec.get("model") or "")
                pricing = get_effective_model_pricing().get(pricing_key) if pricing_key else None
                if pricing:
                    cost_usd += inp * pricing["input"] / 1_000_000
                    cost_usd += out * pricing["output"] / 1_000_000
            elif secs:
                cost_usd += get_audio_cost(rec.get("model") or "", secs)
    return {"input": total_input, "output": total_output, "cost_usd": round(cost_usd, 6)}


def model_row_to_dict(row):
    model = dict(row)
    model["provider_label"] = get_provider_label(model["provider"])
    model["type_label"] = MODEL_TYPES.get(model["model_type"], model["model_type"])
    model["has_api_key"] = bool(get_provider_api_key(model["provider"]))
    return model

def seed_default_ai_models(cursor):
    for model in DEFAULT_AI_MODELS:
        cursor.execute(
            "SELECT id FROM ai_models WHERE provider = ? AND model_type = ? AND model_id = ?",
            (model["provider"], model["model_type"], model["model_id"])
        )
        if cursor.fetchone():
            continue

        cursor.execute(
            """
            INSERT INTO ai_models (provider, model_type, display_name, model_id, enabled, is_default)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                model["provider"],
                model["model_type"],
                model["display_name"],
                model["model_id"],
                model.get("enabled", 1),
                model["is_default"]
            )
        )

    for provider in PROVIDERS:
        for model_type in MODEL_TYPES:
            cursor.execute(
                """
                SELECT id FROM ai_models
                WHERE provider = ? AND model_type = ? AND is_default = 1
                LIMIT 1
                """,
                (provider, model_type)
            )
            if cursor.fetchone():
                continue

            cursor.execute(
                """
                UPDATE ai_models
                SET is_default = 1
                WHERE id = (
                    SELECT id FROM ai_models
                    WHERE provider = ? AND model_type = ? AND enabled = 1
                    ORDER BY id ASC
                    LIMIT 1
                )
                """,
                (provider, model_type)
            )

def list_ai_models(model_type=None, enabled_only=False):
    conn = get_db_connection()
    cursor = conn.cursor()
    query = "SELECT * FROM ai_models"
    filters = []
    params = []

    if model_type:
        filters.append("model_type = ?")
        params.append(model_type)
    if enabled_only:
        filters.append("enabled = 1")

    if filters:
        query += " WHERE " + " AND ".join(filters)

    query += """
        ORDER BY
            model_type DESC,
            provider ASC,
            is_default DESC,
            display_name ASC
    """

    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    return [model_row_to_dict(row) for row in rows]

def get_ai_model_by_id(model_id, model_type=None, enabled_only=True):
    try:
        model_pk = int(model_id)
    except (TypeError, ValueError):
        return None

    conn = get_db_connection()
    cursor = conn.cursor()
    query = "SELECT * FROM ai_models WHERE id = ?"
    params = [model_pk]

    if model_type:
        query += " AND model_type = ?"
        params.append(model_type)
    if enabled_only:
        query += " AND enabled = 1"

    cursor.execute(query, params)
    row = cursor.fetchone()
    conn.close()
    return model_row_to_dict(row) if row else None

def get_default_ai_model(model_type, preferred_provider=None):
    conn = get_db_connection()
    cursor = conn.cursor()

    if preferred_provider:
        cursor.execute(
            """
            SELECT * FROM ai_models
            WHERE model_type = ? AND provider = ? AND enabled = 1
            ORDER BY is_default DESC, id ASC
            LIMIT 1
            """,
            (model_type, preferred_provider)
        )
        row = cursor.fetchone()
        if row:
            conn.close()
            return model_row_to_dict(row)

    cursor.execute(
        """
        SELECT * FROM ai_models
        WHERE model_type = ? AND enabled = 1
        ORDER BY is_default DESC, provider ASC, id ASC
        LIMIT 1
        """,
        (model_type,)
    )
    row = cursor.fetchone()
    conn.close()
    return model_row_to_dict(row) if row else None

def get_selected_transcription_model(cloud_model_id):
    selected_model = get_ai_model_by_id(cloud_model_id, "transcription")
    if selected_model:
        return selected_model
    return get_default_ai_model("transcription")

def list_available_openai_chat_models():
    return [
        model
        for model in list_ai_models(model_type="chat", enabled_only=True)
        if model["provider"] == "openai" and model["has_api_key"]
    ]

def list_available_chat_models():
    models = list_ai_models(model_type="chat", enabled_only=True)
    for model in models:
        model["chat_label"] = describe_chat_answer_model(model)
    return models

def get_default_openai_chat_model():
    openai_chat_models = list_available_openai_chat_models()
    for model in openai_chat_models:
        if model["is_default"]:
            return model
    return openai_chat_models[0] if openai_chat_models else None

def get_selected_openai_chat_model(chat_model_id=None):
    if chat_model_id:
        selected_model = get_ai_model_by_id(chat_model_id, "chat")
        if not selected_model or selected_model["provider"] != "openai" or not selected_model["has_api_key"]:
            raise ValueError("Wybrany model czatu OpenAI nie jest dostępny.")
        return selected_model

    selected_model = get_default_openai_chat_model()
    if selected_model:
        return selected_model

    raise ValueError("Brak dostępnego modelu OpenAI typu chat. Sprawdź ustawienia modeli i klucz OPENAI_API_KEY.")

def promote_first_enabled_model_as_default(cursor, provider, model_type):
    cursor.execute(
        """
        SELECT id FROM ai_models
        WHERE provider = ? AND model_type = ? AND enabled = 1
        ORDER BY id ASC
        LIMIT 1
        """,
        (provider, model_type)
    )
    replacement = cursor.fetchone()
    if not replacement:
        return

    cursor.execute(
        "UPDATE ai_models SET is_default = 1 WHERE id = ?",
        (replacement[0],)
    )

def raise_for_openai_error(response):
    if response.ok:
        return

    try:
        error_payload = response.json()
        message = error_payload.get("error", {}).get("message") or response.text
        code = error_payload.get("error", {}).get("code", "")
    except ValueError:
        message = response.text
        code = ""

    if response.status_code == 401 or code == "invalid_api_key":
        raise_invalid_api_key_error("openai")

    raise RuntimeError(f"Błąd OpenAI API ({response.status_code}): {message}")

def transcribe_with_cloud(model_config, file_path, language):
    provider = model_config["provider"]
    api_key = require_provider_api_key(provider)

    if provider == "groq":
        client = Groq(api_key=api_key)
        try:
            with open(file_path, "rb") as audio_file:
                transcription_options = {
                    "file": (audio_file.name, audio_file.read()),
                    "model": model_config["model_id"]
                }
                if language != "auto":
                    transcription_options["language"] = language

                transcription = client.audio.transcriptions.create(**transcription_options)

            duration = get_audio_duration_seconds(file_path)
            if duration is not None:
                usage_record = {
                    "datetime": datetime.now(timezone.utc).isoformat(),
                    "operation": "transcription",
                    "model": model_config["model_id"],
                    "input_tokens": None,
                    "output_tokens": None,
                    "total_tokens": None,
                    "cached_tokens": None,
                    "reasoning_tokens": None,
                    "usage_type": "audio",
                    "seconds": duration,
                    "usage": {"seconds": duration},
                }
                append_openai_usage_to_request(usage_record)
                try:
                    append_openai_usage_to_file(usage_record)
                except Exception:
                    pass

            return transcription.text
        except Exception as error:
            raise_provider_api_error(provider, error)

    if provider == "openai":
        model_id = model_config["model_id"]
        with open(file_path, "rb") as audio_file:
            files = {
                "file": (os.path.basename(file_path), audio_file)
            }
            data = {"model": model_id}
            if language != "auto":
                data["language"] = language
            # whisper-1 is billed per minute; verbose_json includes the duration field
            if model_id == "whisper-1":
                data["response_format"] = "verbose_json"

            response = requests.post(
                "https://api.openai.com/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {api_key}"},
                data=data,
                files=files,
                timeout=180
            )
            raise_for_openai_error(response)
            payload = response.json()
            payload.setdefault("model", model_id)
            record_openai_usage(payload, "transcription")
            return payload.get("text", "")

    raise RuntimeError(f"Nieobsługiwany provider: {provider}")

def chat_with_cloud(messages, model_config):
    if not model_config:
        raise RuntimeError("Brak aktywnego modelu AI w ustawieniach")

    provider = model_config["provider"]

    if provider == "ollama_ip":
        base_url = get_provider_api_key("ollama_ip").rstrip("/")
        if not base_url:
            raise RuntimeError("Brak OLLAMA_IP_URL w pliku .env. Ustaw np. OLLAMA_IP_URL=http://192.168.1.248:11434")
        try:
            response = requests.post(
                f"{base_url}/v1/chat/completions",
                headers={"Content-Type": "application/json"},
                json={"model": model_config["model_id"], "messages": messages},
                timeout=180
            )
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
        except requests.exceptions.RequestException as error:
            raise RuntimeError(f"Błąd połączenia z Ollama ({base_url}): {error}")

    api_key = require_provider_api_key(provider)

    if provider == "groq":
        client = Groq(api_key=api_key)
        try:
            completion = client.chat.completions.create(
                model=model_config["model_id"],
                messages=messages
            )
            return completion.choices[0].message.content
        except Exception as error:
            raise_provider_api_error(provider, error)

    if provider == "openai" and model_config.get("model_type") == "notes":
        return chat_with_openai_responses(messages, model_config)

    if provider == "openai":
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": model_config["model_id"],
                "messages": messages
            },
            timeout=180
        )
        raise_for_openai_error(response)
        payload = response.json()
        record_openai_usage(payload, "analysis")
        return payload["choices"][0]["message"]["content"]

    raise RuntimeError(f"Nieobsługiwany provider: {provider}")

def extract_responses_output_text(payload):
    if payload.get("output_text"):
        return payload["output_text"]

    chunks = []
    for item in payload.get("output", []):
        if item.get("type") != "message":
            continue
        for content in item.get("content", []):
            text = content.get("text")
            if text:
                chunks.append(text)

    return "\n".join(chunks).strip()

def extract_responses_sources(payload):
    sources = []
    seen_urls = set()

    def add_source(url, title=None):
        if not url or url in seen_urls:
            return
        seen_urls.add(url)
        sources.append({
            "url": url,
            "title": title or url
        })

    for item in payload.get("output", []):
        if item.get("type") == "web_search_call":
            action = item.get("action") or {}
            for source in action.get("sources") or []:
                add_source(source.get("url"), source.get("title"))

        if item.get("type") == "message":
            for content in item.get("content", []):
                for annotation in content.get("annotations", []) or []:
                    if annotation.get("type") == "url_citation":
                        add_source(annotation.get("url"), annotation.get("title"))

    return sources

def messages_to_responses_input(messages):
    role_labels = {
        "system": "Instrukcje systemowe",
        "user": "Użytkownik",
        "assistant": "Asystent"
    }
    prompt_parts = []
    for message in messages:
        role = role_labels.get(message.get("role"), message.get("role", "Wiadomość"))
        prompt_parts.append(f"{role}:\n{message.get('content', '')}")

    return "\n\n".join(prompt_parts)

def create_openai_responses(model_id, input_text, tools=None, include=None, timeout=180, operation_name="analysis"):
    api_key = require_provider_api_key("openai")
    payload = {
        "model": model_id,
        "input": input_text
    }
    if tools:
        payload["tools"] = tools
    if include:
        payload["include"] = include
    if tools:
        payload["tool_choice"] = "auto"

    response = requests.post(
        "https://api.openai.com/v1/responses",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        json=payload,
        timeout=timeout
    )
    raise_for_openai_error(response)
    payload = response.json()
    record_openai_usage(payload, operation_name)
    return payload

def chat_with_openai_responses(messages, model_config):
    payload = create_openai_responses(
        model_config["model_id"],
        messages_to_responses_input(messages),
        timeout=180,
        operation_name="analysis"
    )
    output_text = extract_responses_output_text(payload)
    if not output_text:
        raise RuntimeError("OpenAI Responses API nie zwróciło tekstu odpowiedzi.")

    return output_text

def build_web_search_question_prompt(transcription, context_rows, question):
    # print(f"[build_web_search_question_prompt] START => transcription: {transcription[:100]} ....===.... {transcription[:100]}\ncontext_rows: {context_rows}\nquestion: {question}")
    # print(f"[build_web_search_question_prompt] START => transcription len: {len(transcription)}\ncontext_rows: {context_rows}\nquestion: {question}")
    previous_messages = []
    for role, content in context_rows:
        speaker = "Użytkownik" if role == "user" else "Asystent"
        previous_messages.append(f"{speaker}: {content}")

    history_text = "\n".join(previous_messages) if previous_messages else "Brak wcześniejszej rozmowy."

    return (
        "Jesteś asystentem odpowiadającym na pytania. Odpowiadaj po polsku.\n"
        "Masz dostęp do narzędzia web_search. Używaj go wtedy, gdy pytanie wymaga aktualnych, "
        "zewnętrznych lub weryfikowalnych informacji. Jeśli odpowiedź wynika "
        "wyłącznie z transkrypcji, oprzyj się na transkrypcji i jasno to zaznacz.\n"
        "Nie zmyślaj faktów. Nie dodawaj własnych tez ani teorii. Opieraj sie na dowodach, faktach i stanach faktycznych.\n"
        "Gdy korzystasz z internetu, wskaż źródła i podaj: domenę / autora / date publikacji jeżeli dostępne.\n\n"
        "Kontekst poprzedniej rozmowy:\n"
        f"{history_text}\n\n"
        + (f"Transkrypcja nagrania:\n{transcription}\n\n" if transcription else "")
        + f"Pytanie użytkownika:\n{question}"
    )

def analyze_with_web_search(response_prompt, model_config=None):
    model_id = model_config["model_id"] if model_config else OPENAI_WEB_SEARCH_MODEL
    payload = create_openai_responses(
        model_id,
        response_prompt,
        tools=[
            {
                "type": "web_search"
            }
        ],
        include=["web_search_call.action.sources"],
        timeout=240,
        operation_name="web_search_analysis"
    )

    output_text = extract_responses_output_text(payload)
    if not output_text:
        raise RuntimeError("OpenAI Responses API nie zwróciło tekstu odpowiedzi.")

    return {
        "text": output_text,
        "sources": extract_responses_sources(payload),
        "model": model_id,
        "engine": f"{describe_cloud_model(model_config)} + web_search" if model_config else f"OpenAI Responses API ({model_id}) + web_search"
    }

def answer_question_with_ai(response_prompt, model_config):
    # print(f"[answer_question_with_ai] START => response_prompt len: {len(response_prompt)}, model_config: {model_config}")
    if not model_config:
        raise RuntimeError("Brak aktywnego modelu czatu w ustawieniach")

    if model_config["provider"] == "openai":
        return analyze_with_web_search(response_prompt, model_config)

    answer_text = chat_with_cloud([{"role": "user", "content": response_prompt}], model_config)
    return {
        "text": answer_text,
        "sources": [],
        "model": model_config["model_id"],
        "engine": describe_chat_answer_model(model_config)
    }

def is_port_available(host, port):
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        try:
            sock.bind((host, port))
        except OSError:
            return False
        return True

def find_available_port(start_port, host='0.0.0.0', max_attempts=50):
    for port in range(start_port, start_port + max_attempts):
        if is_port_available(host, port):
            return port

    last_port = start_port + max_attempts - 1
    raise RuntimeError(f"Brak wolnego portu w zakresie {start_port}-{last_port}")

def get_start_port(default_port=8000):
    raw_port = os.getenv('PORT', str(default_port)).strip()
    try:
        return int(raw_port)
    except ValueError:
        print(f"Niepoprawna wartość PORT={raw_port}. Używam portu {default_port}.")
        return default_port

def get_server_port(host='0.0.0.0', default_port=8000):
    selected_port = os.getenv('APP_SELECTED_PORT')
    if selected_port:
        try:
            return int(selected_port)
        except ValueError:
            os.environ.pop('APP_SELECTED_PORT', None)

    start_port = get_start_port(default_port)
    port = find_available_port(start_port, host)
    os.environ['APP_SELECTED_PORT'] = str(port)

    if port != start_port:
        print(f"Port {start_port} jest zajęty. Uruchamiam aplikację na porcie {port}.")

    return port

def find_deno_runtime():
    candidate_paths = [
        os.getenv('YT_DLP_DENO_PATH', '').strip(),
        shutil.which('deno'),
        os.path.expanduser('~/.deno/bin/deno')
    ]

    for path in candidate_paths:
        if path and os.path.isfile(path) and os.access(path, os.X_OK):
            return path

    return None

def build_youtube_download_options(download_token):
    ydl_opts = {
        'cookiefile': '/opt/apps/app_projectstt/data_files/cookies.txt',
        'format': 'bestaudio/best',
        'outtmpl': os.path.join(UPLOAD_FOLDER, f'{download_token}_%(id)s.%(ext)s'),
        'quiet': True,
        'noplaylist': True,
        'overwrites': True,
        # iOS client bypasses n-challenge entirely; web_creator as secondary fallback
        'extractor_args': {
            'youtube': {
                'player_client': ['ios', 'web_creator'],
            }
        },
        # download EJS solver from GitHub if n-challenge is still needed
        'remote_components': 'ejs:github',
    }

    deno_path = find_deno_runtime()
    if deno_path:
        ydl_opts['js_runtimes'] = {'deno': {'path': deno_path}}

    return ydl_opts

def make_temp_upload_path(filename):
    safe_name = secure_filename(filename) or "upload"
    return os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}_{safe_name}")

def resolve_youtube_download_path(ydl, info, download_token):
    candidates = []

    for download in info.get("requested_downloads") or []:
        for key in ("filepath", "filename"):
            p = download.get(key)
            if p:
                candidates.append(p)

    prepared_path = ydl.prepare_filename(info)
    if prepared_path:
        candidates.append(prepared_path)

    video_id = info.get("id")
    if video_id:
        candidates.extend(glob.glob(os.path.join(UPLOAD_FOLDER, f"{download_token}_{video_id}.*")))

    candidates.extend(glob.glob(os.path.join(UPLOAD_FOLDER, f"{download_token}_*")))

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate

    raise FileNotFoundError(
        f"Nie znaleziono pobranego pliku YouTube dla tokenu {download_token}. "
        f"Szukano w: {candidates}"
    )

def _assert_safe_url(url):
    from urllib.parse import urlparse
    parsed = urlparse(url)
    if parsed.scheme not in ('http', 'https'):
        raise ValueError(f"Niedozwolony protokół URL: '{parsed.scheme}'")
    hostname = parsed.hostname
    if not hostname:
        raise ValueError("Brak nazwy hosta w URL")
    try:
        resolved_ip = socket.gethostbyname(hostname)
        ip = ipaddress.ip_address(resolved_ip)
    except socket.gaierror:
        raise ValueError(f"Nie można rozwiązać nazwy hosta: {hostname}")
    if any(ip in net for net in _PRIVATE_NETWORKS) or ip.is_loopback or ip.is_reserved:
        raise ValueError(f"Adres URL wskazuje na wewnętrzną sieć — niedozwolone")

def _assert_youtube_url(url):
    if not _YOUTUBE_RE.match(url):
        raise ValueError("URL nie jest poprawnym adresem YouTube (youtube.com lub youtu.be)")

def extract_youtube_video_id(url):
    m = re.search(
        r'(?:youtube\.com/(?:watch\?.*v=|shorts/|embed/|live/)|youtu\.be/)([a-zA-Z0-9_-]{11})',
        url
    )
    if not m:
        raise ValueError(f"Nie można wyodrębnić ID wideo z URL: {url}")
    return m.group(1)


def fetch_youtube_transcript(url):
    """Pobiera napisy przez YouTube Transcript API (v1.x). Szybkie, nie wymaga pobierania audio."""
    video_id = extract_youtube_video_id(url)
    snippets = _ytt_api.fetch(video_id, languages=['pl', 'en'])
    text = ' '.join(s.text for s in snippets).strip()
    if not text:
        raise ValueError("Napisy są puste")
    return {
        "text": text,
        "title": f"YT: {video_id}",
        "video_id": video_id,
        "source": "transcript_api",
    }

def download_youtube_audio(youtube_url):
    _assert_youtube_url(youtube_url)
    download_token = uuid.uuid4().hex
    ydl_opts = build_youtube_download_options(download_token)

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(youtube_url, download=True)
        return {
            "file_path": resolve_youtube_download_path(ydl, info, download_token),
            "title": info.get("title", "YouTube Video"),
            "id": info.get("id"),
            "duration": info.get("duration"),
            "webpage_url": info.get("webpage_url") or youtube_url
        }

def _is_multi_url_text(text):
    """True gdy każda niepusta linia zaczyna się od http:// lub https://"""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return len(lines) >= 1 and all(l.startswith('http://') or l.startswith('https://') for l in lines)

def fetch_multi_url_content(text):
    """Pobiera treść każdego URL z tekstu i łączy wyniki."""
    urls = [l.strip() for l in text.splitlines() if l.strip()]
    parts = []
    titles = []
    for url in urls:
        try:
            content, title = fetch_webpage_content(url)
            parts.append(f"=== Źródło: {title or url} ===\n{content}")
            titles.append(title or url)
        except Exception as exc:
            app.logger.warning("[multi-url] Błąd pobierania %s: %s", url, exc)
            parts.append(f"=== Błąd pobierania {url}: {exc} ===")
    combined = "\n\n".join(parts)
    display = f"Synteza {len(urls)} źródeł" if len(urls) > 1 else (titles[0] if titles else urls[0])
    return combined, display

def fetch_webpage_content(url):
    _assert_safe_url(url)
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; projectSTT/1.0)'}
    response = requests.get(url, timeout=30, headers=headers)
    response.raise_for_status()

    tree = lxml_html.fromstring(response.content)

    title_nodes = tree.xpath('//title/text()')
    page_title = title_nodes[0].strip() if title_nodes else ""

    for tag in ('script', 'style', 'nav', 'footer', 'header', 'aside', 'noscript'):
        for el in tree.xpath(f'//{tag}'):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    raw_text = tree.text_content()
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    return '\n'.join(lines), page_title

# Funkcja budująca prompt dla generowania notatek AI z transkrypcji audio
def build_audio_notes_prompt(raw_text, mode='full'):
    if mode == 'reel-prepare':
        # ── TUTAJ WPISZ SWÓJ PROMPT DLA TRYBU "PRZYGOTUJ ROLKĘ" ──────────────
        # Przykład struktury:
        #   return (
        #       "Jesteś twórcą treści wideo. Na podstawie poniższego tekstu...\n\n"
        #       f"Tekst:\n{raw_text}"
        #   )
        # ─────────────────────────────────────────────────────────────────────

        # phase=content → tylko generacja treści (short_text) + zapis na S3
        # phase=audio → tylko audio/SRT (wywołanie /api/news-to-video-generator z audio_only=true)
        # phase=shotstack_stage → tylko payload do stage
        # phase=shotstack_prod → tylko payload do prod
        # phase=all → pełny pipeline

        # raw_text_prompt = "Na podstawie dostarczonych materiałów, przygotuj zwięzłą i angażującą rolkę wideo (Reel) na YouTube lub wskazaną platformę.\n"
        # "Rolka powinna mieć do 30 sekund długości lub do wskazanej w transkrypcji długości i zawierać najważniejsze, najbardziej interesujące lub zaskakujące informacje z tekstu. \n"
        # "Skup się na tym, co może przyciągnąć uwagę widzów. Na konic zachęć ich do dalszego zgłębiania tematu. Nie ograniczaj się do prostego streszczenia — poszukaj unikalnych, intrygujących fragmentów, które wyróżniają się w tekście.\n"
        # "Odpowiedz wyłącznie gotowym scenariuszem rolki, bez żadnych dodatkowych wstępów, komentarzy czy wyjaśnień. Użyj języka polskiego i formatu, który łatwo można przekształcić w dynamiczną, wizualnie atrakcyjną rolkę wideo.\n"
        # "Dołącz do textu adresy url możliwych do wykorzystania obrazów lub grafik pochodzących z dostarczonych źródeł.\n"
        return (
            "Na podstawie poniższego tekstu przygotuj payload JSON do generowania rolki wideo.\n"
            "Dla treści z domeny londynek.net wydobądź section i id z URL artykułu.\n"
            "Przykład mapowania URL `https://londynek.net/wiadomosci/article?jdnews_id=111425`: section=wiadomosci, id=111425.\n"
            "Dozwolone wartości section: wiadomosc, czytelnia, newslajt, poradnik, ogloszenia, wydarzenia.\n"
            "Domyślny rozmiar: 9x16. Domyślna phase: shotstack_stage.\n\n"
            "Zwróć WYŁĄCZNIE poprawny JSON (bez żadnego opisu, komentarza ani markdown), w formacie:\n"
            "{\n"
            "  \"section\": \"<section>\",\n"
            "  \"id\": \"<id_artykulu>\",\n"
            "  \"size\": \"9x16\",\n"
            "  \"email\": \"reels@marzec.eu\",\n"
            "  \"voice\": \"Microsoft Server Speech Text to Speech Voice (pl-PL, AgnieszkaNeural)\",\n"
            "  \"provider\": \"microsoft\",\n"
            "  \"style_preset\": \"standard_9x16\",\n"
            "  \"app_mode\": \"dev\",\n"
            "  \"phase\": \"shotstack_stage\"\n"
            "}\n\n"
            f"Tekst:\n{raw_text}"
        )
    if mode == 'prompt':
        return (
            "Jesteś ekspertem w konstruowaniu promptów dla dużych modeli językowych. "
            "Przeczytaj uważnie poniższy tekst i na jego podstawie wygeneruj gotowy, "
            "wysokiej jakości prompt w języku polskim, który użytkownik może wkleić bezpośrednio do dowolnego modelu AI.\n\n"
            "Wygenerowany prompt MUSI:\n"
            "1. Zawierać **precyzyjnie sformułowane pytanie lub zadanie** wynikające z treści tekstu — "
            "tak, aby model AI wiedział dokładnie, czego oczekujesz.\n"
            "2. Zawierać sekcję **KONTEKST** z najważniejszymi informacjami wyciągniętymi z tekstu "
            "(minimum faktów niezbędnych do udzielenia wartościowej odpowiedzi).\n"
            "3. Zawierać sekcję **MUST-HAVE ŹRÓDŁA** — listę konkretnych typów źródeł lub serwisów, "
            "które model POWINIEN przeszukać lub uwzględnić, aby odpowiedź była wiarygodna i aktualna. "
            "Dobierz źródła odpowiednio do tematu tekstu (np. bazy naukowe, serwisy informacyjne, "
            "oficjalne strony instytucji, repozytoria kodu, encyklopedie, fora branżowe itp.).\n"
            "4. Być sformatowany czytelnie dla modelu: używaj nagłówków Markdown (###), "
            "list punktowanych i pogrubień tam, gdzie zwiększają przejrzystość.\n\n"
            "Zwróć WYŁĄCZNIE gotowy prompt — bez żadnego wstępu, komentarza ani wyjaśnienia z Twojej strony.\n\n"
            f"Tekst źródłowy:\n{raw_text}"
        )
    if mode == 'summary':
        return (
            "Jesteś profesjonalnym asystentem. Przeczytaj poniższy tekst "
            "i napisz KRÓTKIE STRESZCZENIE w języku polskim (maksymalnie 5 zdań). "
            "Skup się wyłącznie na najważniejszych informacjach. "
            "Nie stosuj nagłówków ani list.\n\n"
            f"Tekst:\n{raw_text}"
        )
    if mode == 'overview':
        return (
            "Jesteś profesjonalnym asystentem. Przeczytaj poniższy tekst "
            "i napisz KRÓTKIE OMÓWIENIE w języku polskim. "
            "Omów główne tematy, wnioski i kontekst w 2-4 czytelnych akapitach. "
            "Nie stosuj list punktowanych ani sekcji zadań.\n\n"
            f"Tekst:\n{raw_text}"
        )
    if mode == 'bullets':
        return (
            "Jesteś profesjonalnym asystentem. Przeczytaj poniższy tekst "
            "i wypisz LISTĘ NAJWAŻNIEJSZYCH PUNKTÓW w języku polskim. "
            "Każdy punkt zacznij od myślnika (-). Wypisz od 5 do 10 kluczowych punktów. "
            "Odpowiedz TYLKO listą — bez wstępu i komentarzy.\n\n"
            f"Tekst:\n{raw_text}"
        )
    # mode == 'full' — domyślny pełny prompt
    # print(f"[build_audio_notes_prompt] START => raw_text len: {len(raw_text)}")

    how_to_use_web_search_prompt = """Jeśli odpowiedź może być udzielona na podstawie transkrypcji i nie wymaga użycia źródeł zewnętrznych, jasno to zaznacz.
Nie zmyślaj faktów. Nie dodawaj własnych tez ani teorii. Opieraj sie na dowodach, faktach i stanach faktycznych.
Gdy korzystasz z internetu, wskaż źródła wiedzy, podając: domenę / autora / date publikacji (jeżeli dostępna).
Jeżeli dla potwierdzenia autentyczności tekstu konieczne jest użycie web_search, przygotuj dokładny prompt dla modelu weryfikacyjnego, 
opisujący jakie frazy i źródła należy sprawdzić, aby potwierdzić autentyczność badanego tekstu. 
Możliwe źródła video do przeanalizowania to: YouTube, Twitter, Facebook, Instagram, TikTok, Vimeo, Dailymotion. 
Możliwe źródła tekstowe to: Google Search, Bing Search, DuckDuckGo Search, Wikipedia, news media, blogi, fora internetowe.
"""
    model_info_web_search_prompt = """Tutaj napisz prompt dla modelu dotyczący analizy textu przy użyciu web_search. 
Opisz w nim dokładnie, jakie frazy i źródła należy sprawdzić, aby potwierdzić autentyczność badanego tekstu. 
Pamiętaj, że możliwe źródła video do przeanalizowania to: YouTube, Twitter, Facebook, Instagram, TikTok, Vimeo, Dailymotion. 
Możliwe źródła tekstowe to: Google Search, Bing Search, DuckDuckGo Search, Wikipedia, news media, blogi, fora internetowe. 
Na końcu udzielonej odpowiedzi wypisz adresy url, na podstawie których została zweryfikowana autentyczność tekstu."""
    model_info_web_search_transc = "Tutaj wstaw skrócony opis transkrypcji do weryfikacji. Nie wstawaj całej transkrypcji, tylko jej streszczenie, które pozwoli zrozumieć, o czym jest tekst i jakie informacje zawiera."
    verification_json_format = "({\"prompt\": \"%s\", \"transcription\": \"%s\"})", model_info_web_search_prompt, model_info_web_search_transc

    return (
        "Jesteś profesjonalnym asystentem biurowym. Przeczytaj uważnie tekst do przeanalizowania "
        "i przygotuj z niego czytelną, ustrukturyzowaną notatkę w języku polskim.\n"
        "Notatka MUSI składać się z trzech wyraźnych sekcji, w poniższej kolejności:\n"
        "### KRÓTKIE PODSUMOWANIE\n"
        "Na każde, około 100 zdań treści, użyj 2-3 zdania podsumowania opisujące esencję przeanalizowanej treści.\n\n"
        "### NAJWAŻNIEJSZE PUNKTY\n"
        "Kluczowe informacje i wątki wypisane od myślników.\n\n"
        "### LISTA ZADAŃ DO WYKONANIA\n"
        "Zadania i akcje do podjęcia (jeśli wspomniano); jeśli brak — napisz: Brak zadań.\n\n"
        "### PROMPT DO SPRAWDZENIA AUTENTYCZNOŚCI\n"
        "Na podstawie treści napisz gotowy prompt weryfikacyjny w poniższym formacie JSON:\n"
        f"{verification_json_format}\n\n"
        f"PAMIĘTAJ: {how_to_use_web_search_prompt}\n"
        "### WYNIK AUTENTYCZNOŚCI\n"
        "Oceń wiarygodność treści w skali 0-100 i zwróć WYŁĄCZNIE poniższy JSON (bez dodatkowego tekstu):\n"
        "{\"authenticity_score\": N}\n"
        "N=100 oznacza treść w pełni wiarygodną, N=0 oznacza treść fałszywą lub niemożliwą do zweryfikowania.\n\n"
        f"Oto tekst do przeanalizowania:\n{raw_text}"
    )

    # return (
    #     "Jesteś profesjonalnym asystentem biurowym. Przeczytaj uważnie poniższy tekst "
    #     "i przygotuj z niego czytelną, ustrukturyzowaną notatkę w języku polskim.\n"
    #     "Notatka MUSI składać się z trzech wyraźnych sekcji, w poniższej kolejności:\n"
    #     "1. KRÓTKIE PODSUMOWANIE (2-4 zdania wyjaśniające esencję nagrania).\n"
    #     "2. NAJWAŻNIEJSZE PUNKTY (kluczowe informacje i wątki wypisane od myślników).\n"
    #     "3. LISTA ZADAŃ DO WYKONANIA - Zadania i akcje do podjęcia (jeśli wspomniano); jeśli brak — napisz: Brak zadań.\n"
    #     "4. Jeśli treść zawiera twierdzenia faktograficzne wymagające weryfikacji, dodaj sekcję "
    #     "'PROMPT DO SPRAWDZENIA AUTENTYCZNOŚCI'. W tej sekcji opisz, jak przeszukać internet "
    #     "i jakie źródła porównać, aby potwierdzić, czy tekst jest autentyczny i nie jest fake newsem.\n"
    #     "   Sekcję 'PROMPT DO SPRAWDZENIA AUTENTYCZNOŚCI' przedstaw jako poprawny JSON w formacie:\n"
    #     f"{verification_json_format}\n\n"
    #     f"Oto tekst do przeanalizowania:\n{raw_text}"
    # )

    # return (
    #     "Jesteś profesjonalnym asystentem biurowym. Przeczytaj uważnie poniższy tekst "
    #     "i przygotuj z niego czytelną, ustrukturyzowaną notatkę w języku polskim.\n\n"
    #     "Notatka MUSI zawierać cztery sekcje w podanej kolejności:\n\n"
    #     "### KRÓTKIE PODSUMOWANIE\n"
    #     "2-4 zdania opisujące esencję treści.\n\n"
    #     "### NAJWAŻNIEJSZE PUNKTY\n"
    #     "Kluczowe informacje i wątki wypisane od myślników.\n\n"
    #     "### LISTA ZADAŃ DO WYKONANIA\n"
    #     "Zadania i akcje do podjęcia (jeśli wspomniano); jeśli brak — napisz: Brak zadań.\n\n"
    #     "### PROMPT DO SPRAWDZENIA AUTENTYCZNOŚCI\n"
    #     "Masz dostęp do narzędzia web_search. Używaj go wtedy, gdy pytanie wymaga aktualnych, "
    #     "zewnętrznych lub weryfikowalnych informacji spoza transkrypcji. Jeśli odpowiedź wynika "
    #     "wyłącznie z transkrypcji, oprzyj się na transkrypcji i jasno to zaznacz.\n"
    #     "Nie zmyślaj faktów. Nie dodawaj własnych tez ani teorii. Opieraj sie na dowodach, faktach i stanach faktycznych.\n"
    #     "Gdy korzystasz z internetu, wskaż źródła i podaj: domenę / autora / date publikacji jeżeli dostępne.\n\n"
    #     "Na podstawie treści napisz gotowy prompt weryfikacyjny w formacie JSON:\n"
    #     '{"prompt": "<tu wpisz konkretne zapytanie weryfikacyjne dla analizy treści uwzględniające: '
    #     "(1) frazy i źródła tekstowe do sprawdzenia, "
    #     "(2) zapytania do YouTube, wyszukiwarek oraz platform o wideo powiązane z tematem, "
    #     "(3) Porównaj informacje na stronach rządowych, zdania ekspertów, tematyczne fora internetowe oraz lokalne przepisy."
    #     '(4) instrukcję jak porównać źródła i ocenić autentyczność>"}\n\n'
    #     "WAŻNE: Pole prompt musi być jednotematycznym, konkretnym zapytaniem opartym na treści — "
    #     "nie szablonem ani nie opisem zadania. Cały JSON musi być w jednej linii, bez znaków nowej linii wewnątrz wartości.\n\n"
    #     f"Tekst do przeanalizowania:\n{raw_text}"
    # )

# Funkcja generująca notatki AI na podstawie transkrypcji audio
def extract_authenticity_score(notes_text):
    if not notes_text:
        return None
    match = re.search(r'"authenticity_score"\s*:\s*(\d+)', notes_text)
    if not match:
        return None
    return max(0, min(100, int(match.group(1))))


def generate_audio_notes(raw_text, processing_mode, preferred_provider=None, model_used=None, notes_mode='full', notes_model_id=None):
    if not raw_text.strip():
        return "", None

    prompt = build_audio_notes_prompt(raw_text, mode=notes_mode)
    # print(f"[generate_audio_notes] Built prompt len: {len(prompt)} ==> {prompt}\n\n")

    if processing_mode == "online":
        notes_model = (
            get_ai_model_by_id(notes_model_id, "notes")
            if notes_model_id
            else None
        ) or get_default_ai_model("notes", preferred_provider=preferred_provider)
        return chat_with_cloud([{"role": "user", "content": prompt}], notes_model), notes_model

    try:
        response = ollama.chat(model='llama3', messages=[{'role': 'user', 'content': prompt}])
        return response['message']['content'], None
    except Exception:
        return "Nie udało się wygenerować notatek AI. Upewnij się, że Ollama działa w tle.", None

def normalize_processing_mode(processing_mode):
    mode = str(processing_mode or "offline").strip().lower()
    if mode not in {"offline", "online"}:
        raise ValueError("Nieobsługiwany tryb przetwarzania. Użyj 'offline' albo 'online'.")
    return mode

def get_default_processing_mode():
    return "offline" if LOCAL_MODELS_ENABLED else "online"

# Główna funkcja obsługująca transkrypcję audio i generowanie notatek AI
def process_audio_transcription(file_path, processing_mode='offline', model_name='base', cloud_model_id=None, language='auto', task='transcribe', notes_mode='full', notes_model_id=None):
    processing_mode = normalize_processing_mode(processing_mode)
    model_name = str(model_name or "base").strip()
    language = str(language or "auto").strip()
    task = str(task or "transcribe").strip()

    if task not in {"transcribe", "translate"}:
        raise ValueError("Nieobsługiwane zadanie. Użyj 'transcribe' albo 'translate'.")

    if processing_mode == "online":
        transcription_model = get_selected_transcription_model(cloud_model_id)
        if not transcription_model:
            raise ValueError("Brak aktywnego modelu transkrypcji w ustawieniach")

        raw_text = transcribe_with_cloud(transcription_model, file_path, language)
        notes, notes_model = generate_audio_notes(
            raw_text,
            processing_mode,
            preferred_provider=transcription_model["provider"],
            model_used=str(model_name or cloud_model_id).strip(),
            notes_mode=notes_mode,
            notes_model_id=notes_model_id
        )

        return {
            "text": raw_text,
            "notes": notes,
            "notes_model_used": describe_notes_model(notes_model, processing_mode),
            "language": language if language != "auto" else "auto",
            "task": task,
            "model_used": f"{describe_cloud_model(transcription_model)} + notatki: {describe_cloud_model(notes_model)}"
        }

    if not LOCAL_MODELS_ENABLED:
        raise ValueError("Lokalne modele Whisper są wyłączone. Użyj trybu chmurowego albo uruchom aplikację bez parametru --no-local-models.")

    if model_name not in models:
        raise ValueError("Model lokalny nie jest obsługiwany")

    model = models[model_name]
    options = {"task": task, "fp16": False}
    if language != "auto":
        options["language"] = language

    result = model.transcribe(file_path, **options)
    raw_text = result["text"]
    notes, _ = generate_audio_notes(raw_text, processing_mode, preferred_provider=None, model_used=model_name, notes_mode=notes_mode)

    return {
        "text": raw_text,
        "notes": notes,
        "notes_model_used": describe_notes_model(None, processing_mode),
        "language": result.get("language", language),
        "task": task,
        "model_used": f"Lokalny Whisper ({model_name})"
    }

def save_transcription_history(user_email, display_title, raw_text, notes, notes_model_used="", openai_usage_history=None, project_id=None):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    if project_id is None:
        cursor.execute(
            "INSERT INTO projects (user_email, name) VALUES (?, ?)",
            (user_email, display_title)
        )
        project_id = cursor.lastrowid
    cursor.execute(
        """
        INSERT INTO history (user_email, filename, raw_text, ai_notes, notes_model_used, openai_usage_history, project_id)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            user_email,
            display_title,
            raw_text.strip(),
            notes.strip(),
            str(notes_model_used or "").strip(),
            serialize_openai_usage_history(openai_usage_history),
            project_id
        )
    )
    record_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return record_id, project_id

def get_project_sources(project_id, user_email):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, filename, raw_text FROM history WHERE project_id = ? AND user_email = ? ORDER BY created_at",
        (project_id, user_email)
    )
    rows = cursor.fetchall()
    conn.close()
    return [{"id": r[0], "filename": r[1], "raw_text": r[2]} for r in rows]

def build_project_context(sources, max_chars=60_000):
    parts = []
    total = 0
    for i, src in enumerate(sources, 1):
        header = f"=== Źródło {i}: {src['filename']} ==="
        text = src['raw_text'] or ''
        chunk = f"{header}\n{text}"
        if total + len(chunk) > max_chars:
            remaining = max_chars - total
            if remaining > len(header) + 100:
                parts.append(f"{header}\n{text[:remaining - len(header) - 1]}")
            break
        parts.append(chunk)
        total += len(chunk)
    return "\n\n".join(parts)

def get_payload_setting(payload, settings, key, default=None):
    if key in settings:
        return settings.get(key)
    return payload.get(key, default)

def parse_bool_setting(value, default=False):
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "tak", "on"}
    return bool(value)

def build_local_model_list(model_type=None):
    if not LOCAL_MODELS_ENABLED:
        return []

    requirements = {item["id"]: item for item in LOCAL_MODEL_REQUIREMENTS}
    model_order = {"tiny": 1, "base": 2, "small": 3}
    local_models = []

    if model_type in (None, "transcription"):
        for model_name in sorted(models.keys(), key=lambda name: model_order.get(name, 99)):
            requirement = requirements.get(f"whisper-{model_name}", {})
            local_models.append({
                "source": "local",
                "id": model_name,
                "model_id": model_name,
                "model_type": "transcription",
                "type_label": MODEL_TYPES["transcription"],
                "display_name": requirement.get("name", f"Whisper {model_name.title()}"),
                "engine": requirement.get("engine", "openai-whisper"),
                "enabled": True,
                "available": True,
                "is_default": model_name == "base",
                "summary": requirement.get("summary", ""),
                "request_settings": {
                    "processing_mode": "offline",
                    "model_name": model_name
                }
            })

    for local_chat_type in ("chat", "notes"):
        if model_type not in (None, local_chat_type):
            continue

        requirement = requirements.get("ollama-llama3", {})
        local_models.append({
            "source": "local",
            "id": "ollama-llama3" if local_chat_type == "chat" else "ollama-llama3-notes",
            "model_id": "llama3",
            "model_type": local_chat_type,
            "type_label": MODEL_TYPES[local_chat_type],
            "display_name": requirement.get("name", "Llama 3 przez Ollama"),
            "engine": requirement.get("engine", "ollama"),
            "enabled": True,
            "available": True,
            "is_default": True,
            "summary": requirement.get("summary", ""),
            "request_settings": {
                "processing_mode": "offline",
                "chat_model": "llama3"
            }
        })

    return local_models

def build_cloud_model_list(model_type=None, include_disabled=False):
    cloud_models = []
    for model in list_ai_models(model_type=model_type, enabled_only=not include_disabled):
        enabled = bool(model["enabled"])
        has_api_key = bool(model["has_api_key"])
        cloud_models.append({
            "source": "cloud",
            "id": model["id"],
            "provider": model["provider"],
            "provider_label": model["provider_label"],
            "model_id": model["model_id"],
            "model_type": model["model_type"],
            "type_label": model["type_label"],
            "display_name": model["display_name"],
            "enabled": enabled,
            "available": enabled and has_api_key,
            "is_default": bool(model["is_default"]),
            "request_settings": {
                "processing_mode": "online",
                "cloud_model_id": model["id"]
            }
        })

    return cloud_models

def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            email TEXT PRIMARY KEY,
            first_name TEXT NOT NULL,
            last_name TEXT NOT NULL,
            password_hash TEXT NOT NULL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT NOT NULL,
            name       TEXT NOT NULL DEFAULT 'Nowy projekt',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_email) REFERENCES users(email) ON DELETE CASCADE
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_projects_user
        ON projects(user_email, created_at)
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT NOT NULL,
            filename TEXT NOT NULL,
            raw_text TEXT NOT NULL,
            ai_notes TEXT NOT NULL,
            notes_model_used TEXT DEFAULT '',
            openai_usage_history TEXT DEFAULT '',
            project_id INTEGER REFERENCES projects(id) ON DELETE SET NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_email) REFERENCES users(email)
        )
    ''')
    cursor.execute("PRAGMA table_info(history)")
    history_columns = {row[1] for row in cursor.fetchall()}
    if "notes_model_used" not in history_columns:
        cursor.execute("ALTER TABLE history ADD COLUMN notes_model_used TEXT DEFAULT ''")
    if "openai_usage_history" not in history_columns:
        cursor.execute("ALTER TABLE history ADD COLUMN openai_usage_history TEXT DEFAULT ''")
    if "project_id" not in history_columns:
        cursor.execute("ALTER TABLE history ADD COLUMN project_id INTEGER REFERENCES projects(id) ON DELETE SET NULL")
    # migrate: każdy history bez projektu dostaje własny projekt (single-source)
    cursor.execute("SELECT id, user_email, filename, created_at FROM history WHERE project_id IS NULL")
    for hid, email, fname, hcreated in cursor.fetchall():
        cursor.execute(
            "INSERT INTO projects (user_email, name, created_at) VALUES (?, ?, ?)",
            (email, fname, hcreated)
        )
        cursor.execute("UPDATE history SET project_id = ? WHERE id = ?", (cursor.lastrowid, hid))

    # NOWA TABELA: Pamięć czatu (Prawdziwa rozmowa)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            record_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            model_used TEXT DEFAULT '',
            openai_usage_history TEXT DEFAULT '',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(record_id) REFERENCES history(id) ON DELETE CASCADE
        )
    ''')
    cursor.execute("PRAGMA table_info(chat_history)")
    chat_history_columns = {row[1] for row in cursor.fetchall()}
    if "model_used" not in chat_history_columns:
        cursor.execute("ALTER TABLE chat_history ADD COLUMN model_used TEXT DEFAULT ''")
    if "openai_usage_history" not in chat_history_columns:
        cursor.execute("ALTER TABLE chat_history ADD COLUMN openai_usage_history TEXT DEFAULT ''")

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS project_chat_history (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER NOT NULL,
            role       TEXT NOT NULL,
            content    TEXT NOT NULL,
            model_used TEXT DEFAULT '',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_proj_chat_lookup
        ON project_chat_history(project_id, id)
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS compare_results (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email    TEXT NOT NULL,
            display_name  TEXT NOT NULL,
            provider      TEXT NOT NULL,
            notes_mode    TEXT NOT NULL DEFAULT 'full',
            source_chars  INTEGER NOT NULL DEFAULT 0,
            tokens_in_est INTEGER NOT NULL DEFAULT 0,
            tokens_out_est INTEGER NOT NULL DEFAULT 0,
            timing_ms     INTEGER NOT NULL DEFAULT 0,
            success       INTEGER NOT NULL DEFAULT 1,
            created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_email) REFERENCES users(email) ON DELETE CASCADE
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_compare_results_user
        ON compare_results(user_email, created_at)
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS password_reset_tokens (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            email      TEXT NOT NULL,
            code       TEXT NOT NULL,
            expires_at DATETIME NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_prt_email
        ON password_reset_tokens(email)
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS app_logs (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT,
            level      TEXT NOT NULL DEFAULT 'error',
            category   TEXT NOT NULL DEFAULT 'general',
            message    TEXT NOT NULL,
            detail     TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_app_logs_user
        ON app_logs(user_email, created_at)
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS ai_models (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            provider TEXT NOT NULL,
            model_type TEXT NOT NULL,
            display_name TEXT NOT NULL,
            model_id TEXT NOT NULL,
            enabled INTEGER NOT NULL DEFAULT 1,
            is_default INTEGER NOT NULL DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_ai_models_lookup
        ON ai_models(model_type, provider, enabled, is_default)
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS api_keys (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT NOT NULL,
            key_hash TEXT NOT NULL UNIQUE,
            name TEXT NOT NULL DEFAULT '',
            allowed_origin TEXT NOT NULL DEFAULT '',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            last_used_at DATETIME,
            FOREIGN KEY(user_email) REFERENCES users(email) ON DELETE CASCADE
        )
    ''')
    # migration: add column to existing databases
    try:
        cursor.execute("ALTER TABLE api_keys ADD COLUMN allowed_origin TEXT NOT NULL DEFAULT ''")
    except Exception:
        pass
    seed_default_ai_models(cursor)
    conn.commit()
    conn.close()

init_db()

if LOCAL_MODELS_ENABLED:
    print(f"Ładowanie modeli Whisper... or python app.py --no-local-models")
    models = {
        "tiny": whisper.load_model("tiny"),
        "base": whisper.load_model("base"),
        "small": whisper.load_model("small")
    }
else:
    print("Pominięto ładowanie lokalnych modeli Whisper (--no-local-models).")
    models = {}

# Trasy dla poszczególnych stron aplikacji
@app.route('/', methods=['GET', 'POST'])
@limiter.limit("20 per minute; 5 per second")
def login():
    if 'user_email' in session:
        return redirect(url_for('index_page'))
        
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT email, first_name, last_name, password_hash FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()
        conn.close()
        
        if user and check_password_hash(user[3], password):
            session['user_email'] = user[0]
            flash('Zalogowano pomyślnie!', 'success')
            return redirect(url_for('index_page'))
        else:
            flash('Błędny email lub hasło.', 'danger')
        
    return render_template('login-page.html')

def send_email(to_address, subject, body_text):
    smtp_server   = os.environ.get('SMTP_SERVER', '').strip()
    smtp_port     = int(os.environ.get('SMTP_PORT', '465') or '465')
    smtp_user     = os.environ.get('SMTP_USER', '').strip()
    smtp_password = os.environ.get('SMTP_PASSWORD', '').strip()
    sender_email  = os.environ.get('SENDER_EMAIL', smtp_user).strip() or smtp_user
    if not (smtp_server and smtp_user and smtp_password):
        raise RuntimeError("Brak konfiguracji SMTP (SMTP_SERVER, SMTP_USER, SMTP_PASSWORD)")
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From']    = sender_email
    msg['To']      = to_address
    msg.attach(MIMEText(body_text, 'plain', 'utf-8'))
    ctx = _ssl.create_default_context()
    if smtp_port == 465:
        with smtplib.SMTP_SSL(smtp_server, smtp_port, context=ctx) as srv:
            srv.login(smtp_user, smtp_password)
            srv.sendmail(sender_email, to_address, msg.as_bytes())
    else:
        with smtplib.SMTP(smtp_server, smtp_port) as srv:
            srv.ehlo()
            srv.starttls(context=ctx)
            srv.login(smtp_user, smtp_password)
            srv.sendmail(sender_email, to_address, msg.as_bytes())


def log_app_error(user_email, category, message, detail=None, level='error'):
    """Persist a user-facing error to app_logs. Never raises."""
    try:
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO app_logs (user_email, level, category, message, detail) VALUES (?,?,?,?,?)",
            (user_email, level, category,
             str(message)[:1000], str(detail or '')[:4000])
        )
        conn.commit()
        conn.close()
    except Exception:
        pass


@app.route('/reset-password', methods=['GET', 'POST'])
@limiter.limit("5 per hour", methods=["POST"])
def reset_password_request():
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        if not email:
            flash('Podaj adres e-mail.', 'danger')
            return render_template('reset-password.html')

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT email FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()

        if user:
            cursor.execute("DELETE FROM password_reset_tokens WHERE email = ?", (email,))
            code = ''.join(str(random.randint(0, 9)) for _ in range(6))
            expires_at = datetime.utcnow() + timedelta(minutes=10)
            cursor.execute(
                "INSERT INTO password_reset_tokens (email, code, expires_at) VALUES (?, ?, ?)",
                (email, code, expires_at.strftime('%Y-%m-%d %H:%M:%S'))
            )
            conn.commit()
            conn.close()
            try:
                send_email(
                    email,
                    "Kod weryfikacyjny – reset hasła STT",
                    f"Twój jednorazowy kod weryfikacyjny:\n\n  {code}\n\n"
                    f"Kod jest ważny przez 10 minut.\n\n"
                    f"Jeśli nie prosiłeś(-aś) o reset hasła, zignoruj tę wiadomość."
                )
            except Exception as exc:
                app.logger.error("SMTP error: %s", exc)
                log_app_error(email, 'smtp', 'Błąd wysyłania e-maila z kodem resetowania hasła', str(exc))
                flash('Błąd wysyłania e-maila. Sprawdź konfigurację SMTP.', 'danger')
                return render_template('reset-password.html')
        else:
            conn.close()

        # nie zdradzamy czy email istnieje
        flash('Jeśli podany adres e-mail istnieje w systemie, wysłaliśmy kod weryfikacyjny.', 'success')
        return redirect(url_for('reset_password_verify', email=email))

    return render_template('reset-password.html')


@app.route('/reset-password/verify', methods=['GET', 'POST'])
@limiter.limit("10 per 10 minutes", methods=["POST"])
def reset_password_verify():
    email = (request.args.get('email') or request.form.get('email') or '').strip().lower()

    if request.method == 'POST':
        code_input = ''.join(request.form.get(f'c{i}', '') for i in range(6)).strip()
        if not code_input:
            code_input = request.form.get('code', '').strip()

        if not email or len(code_input) != 6:
            flash('Podaj pełny 6-cyfrowy kod.', 'danger')
            return render_template('reset-verify.html', email=email)

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, code, expires_at FROM password_reset_tokens WHERE email = ? ORDER BY created_at DESC LIMIT 1",
            (email,)
        )
        row = cursor.fetchone()

        if not row:
            conn.close()
            flash('Nieprawidłowy kod lub kod wygasł. Spróbuj ponownie.', 'danger')
            return render_template('reset-verify.html', email=email)

        token_id, stored_code, expires_at_str = row
        expires_at = datetime.strptime(expires_at_str, '%Y-%m-%d %H:%M:%S')

        if datetime.utcnow() > expires_at:
            cursor.execute("DELETE FROM password_reset_tokens WHERE id = ?", (token_id,))
            conn.commit()
            conn.close()
            flash('Kod wygasł. Poproś o nowy.', 'danger')
            return redirect(url_for('reset_password_request'))

        if code_input != stored_code:
            conn.close()
            flash('Nieprawidłowy kod. Sprawdź e-mail i spróbuj ponownie.', 'danger')
            return render_template('reset-verify.html', email=email)

        cursor.execute("DELETE FROM password_reset_tokens WHERE email = ?", (email,))
        conn.commit()
        conn.close()

        session['user_email'] = email
        session['password_reset_flow'] = True
        return redirect(url_for('change_password'))

    return render_template('reset-verify.html', email=email)


@app.route('/change-password', methods=['GET', 'POST'])
def change_password():
    if 'user_email' not in session:
        flash('Musisz się zalogować, aby zmienić hasło.', 'danger')
        return redirect(url_for('login'))

    is_reset_flow = session.get('password_reset_flow', False)

    if request.method == 'POST':
        new_password         = request.form.get('new_password', '')
        confirm_new_password = request.form.get('confirm_new_password', '')
        email                = session['user_email']

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()

        pwd_error = validate_password_strength(new_password)

        if not is_reset_flow:
            current_password = request.form.get('current_password', '')
            if not user or not check_password_hash(user[0], current_password):
                flash('Aktualne hasło jest niepoprawne.', 'danger')
                conn.close()
                return render_template('zmiana-hasla.html', is_reset_flow=False)

        if new_password != confirm_new_password:
            flash('Nowe hasła nie są identyczne.', 'danger')
            conn.close()
        elif pwd_error:
            flash(pwd_error, 'danger')
            conn.close()
        else:
            new_hashed = generate_password_hash(new_password)
            cursor.execute("UPDATE users SET password_hash = ? WHERE email = ?", (new_hashed, email))
            conn.commit()
            conn.close()
            session.pop('password_reset_flow', None)
            flash('Hasło zostało pomyślnie zmienione!', 'success')
            return redirect(url_for('index_page'))

    return render_template('zmiana-hasla.html', is_reset_flow=is_reset_flow)

@app.route('/logs')
def logs_page():
    user_email = get_authenticated_user()
    if not user_email:
        flash('Musisz się zalogować.', 'danger')
        return redirect(url_for('login'))
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, level, category, message, detail, datetime(created_at,'localtime') "
        "FROM app_logs WHERE user_email = ? ORDER BY created_at DESC LIMIT 300",
        (user_email,)
    )
    logs = [{"id": r[0], "level": r[1], "category": r[2],
             "message": r[3], "detail": r[4], "created_at": r[5]}
            for r in cursor.fetchall()]
    conn.close()
    return render_template('logs-page.html', logs=logs)


@app.route('/api/logs', methods=['DELETE'])
@csrf.exempt
def clear_logs():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Unauthorized"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM app_logs WHERE user_email = ?", (user_email,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route('/api/logs/<int:log_id>', methods=['DELETE'])
@csrf.exempt
def delete_log(log_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Unauthorized"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM app_logs WHERE id = ? AND user_email = ?", (log_id, user_email))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route('/api/client-log', methods=['POST'])
@limiter.limit("30 per hour")
@csrf.exempt
def client_log():
    """Accept client-side error reports from the frontend."""
    user_email = get_authenticated_user()
    data = request.get_json(silent=True) or {}
    message = str(data.get('message') or '').strip()[:500]
    detail  = str(data.get('detail')  or '').strip()[:2000]
    category = str(data.get('category') or 'frontend').strip()[:50]
    if not message:
        return jsonify({"ok": False}), 400
    log_app_error(user_email, category, message, detail or None)
    return jsonify({"ok": True})


@app.route('/logout')
def logout():
    session.pop('user_email', None)
    flash('Wylogowano pomyślnie.', 'success')
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("10 per hour; 3 per minute")
def register():
    if 'user_email' in session:
        return redirect(url_for('index_page'))

    if request.method == 'POST':
        email = request.form.get('email')
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        password = request.form.get('password')
        
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_regex, email):
            flash('Podaj poprawny adres e-mail (np. nazwa@domena.pl).', 'danger')
            return render_template('rejestracja.html')
            
        pwd_error = validate_password_strength(password)
        if pwd_error:
            flash(pwd_error, 'danger')
            return render_template('rejestracja.html')
        
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT email FROM users WHERE email = ?", (email,))
        existing_user = cursor.fetchone()
        
        if existing_user:
            flash('Ten adres email jest już zarejestrowany!', 'danger')
            conn.close()
        else:
            hashed_password = generate_password_hash(password)
            cursor.execute(
                "INSERT INTO users (email, first_name, last_name, password_hash) VALUES (?, ?, ?, ?)",
                (email, first_name, last_name, hashed_password)
            )
            conn.commit()
            conn.close()
            flash('Rejestracja zakończona sukcesem! Możesz się zalogować.', 'success')
            return redirect(url_for('login'))
        
    return render_template('rejestracja.html')

@app.route('/index')
def index_page():
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))
        
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT email, first_name, last_name FROM users WHERE email = ?", (session['user_email'],))
    user = cursor.fetchone()
    conn.close()
    
    user_data = {
        'email': user[0],
        'first_name': user[1],
        'last_name': user[2]
    }
    cloud_models = list_ai_models(model_type="transcription", enabled_only=True)
    default_cloud_model = get_default_ai_model("transcription")
    default_cloud_model_id = default_cloud_model["id"] if default_cloud_model else None
    default_notes_model = get_default_ai_model("notes")
    openai_chat_models = list_available_openai_chat_models()
    for model in openai_chat_models:
        model["chat_label"] = describe_chat_answer_model(model)
    default_chat_model = get_default_openai_chat_model()
    initial_notes_model_label = (
        describe_notes_model(default_notes_model, "online")
        if not LOCAL_MODELS_ENABLED
        else describe_local_notes_model()
    )
    return render_template(
        'index.html',
        user=user_data,
        cloud_models=cloud_models,
        default_cloud_model_id=default_cloud_model_id,
        initial_notes_model_label=initial_notes_model_label,
        default_cloud_notes_model_label=describe_notes_model(default_notes_model, "online"),
        local_notes_model_label=describe_local_notes_model(),
        openai_chat_models=openai_chat_models,
        default_chat_model_id=default_chat_model["id"] if default_chat_model else None,
        default_chat_model_label=describe_chat_answer_model(default_chat_model) if default_chat_model else "Brak dostępnego modelu OpenAI chat",
        local_models_enabled=LOCAL_MODELS_ENABLED
    )

@app.route('/mobile')
def mobile_page():
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT email, first_name, last_name FROM users WHERE email = ?", (session['user_email'],))
    user = cursor.fetchone()
    conn.close()

    user_data = {
        'email': user[0],
        'first_name': user[1],
        'last_name': user[2]
    }
    cloud_models = list_ai_models(model_type="transcription", enabled_only=True)
    default_cloud_model = get_default_ai_model("transcription")
    default_cloud_model_id = default_cloud_model["id"] if default_cloud_model else None
    default_notes_model = get_default_ai_model("notes")
    all_chat_models = list_available_chat_models()
    openai_chat_models = [m for m in all_chat_models if m["provider"] == "openai"]
    default_chat_model = get_default_ai_model("chat")
    notes_models = list_ai_models(model_type="notes", enabled_only=True)
    default_notes_model_id = default_notes_model["id"] if default_notes_model else None
    initial_notes_model_label = (
        describe_notes_model(default_notes_model, "online")
        if not LOCAL_MODELS_ENABLED
        else describe_local_notes_model()
    )
    ollama_ip_url = get_provider_api_key("ollama_ip")
    return render_template(
        'mobile-page.html',
        user=user_data,
        cloud_models=cloud_models,
        default_cloud_model_id=default_cloud_model_id,
        initial_notes_model_label=initial_notes_model_label,
        default_cloud_notes_model_label=describe_notes_model(default_notes_model, "online"),
        local_notes_model_label=describe_local_notes_model(),
        all_chat_models=all_chat_models,
        openai_chat_models=openai_chat_models,
        default_chat_model_id=default_chat_model["id"] if default_chat_model else None,
        default_chat_model_label=describe_chat_answer_model(default_chat_model) if default_chat_model else "Brak dostępnego modelu czatu",
        notes_models=notes_models,
        default_notes_model_id=default_notes_model_id,
        ollama_ip_configured=bool(ollama_ip_url),
        local_models_enabled=LOCAL_MODELS_ENABLED
    )

@app.route('/usage-history', methods=['GET'])
def usage_history():
    if 'user_email' not in session:
        return redirect(url_for('login'))

    session_email = session['user_email']
    is_admin      = session_email == ADMIN_EMAIL

    # Admin can view any user's stats via ?view=email, or all users via ?view=__all__
    view_email = request.args.get('view', '').strip() or session_email
    if not is_admin:
        view_email = session_email  # non-admins always see themselves

    def _eb():
        return {"input": 0, "output": 0, "cost": 0.0}

    def _fetch_records(conn, email_filter=None):
        c = conn.cursor()
        if email_filter:
            c.execute(
                "SELECT openai_usage_history FROM history "
                "WHERE user_email = ? AND openai_usage_history != ''",
                (email_filter,)
            )
        else:
            c.execute("SELECT openai_usage_history FROM history WHERE openai_usage_history != ''")
        raws = c.fetchall()
        if email_filter:
            c.execute(
                """SELECT ch.openai_usage_history FROM chat_history ch
                   JOIN history h ON ch.record_id = h.id
                   WHERE h.user_email = ? AND ch.openai_usage_history != ''""",
                (email_filter,)
            )
        else:
            c.execute(
                """SELECT ch.openai_usage_history FROM chat_history ch
                   JOIN history h ON ch.record_id = h.id
                   WHERE ch.openai_usage_history != ''"""
            )
        raws += c.fetchall()
        records = []
        for (raw,) in raws:
            records.extend(parse_openai_usage_history(raw))
        return records

    def _aggregate(records):
        today_d = datetime.now(timezone.utc).date()
        totals = _eb(); today_b = _eb(); last7 = _eb(); last30 = _eb()
        by_day = {}; by_model = {}; by_op = {}

        for rec in records:
            inp  = int(rec.get("input_tokens")  or 0)
            out  = int(rec.get("output_tokens") or 0)
            secs = float(rec.get("seconds") or 0)
            model = (rec.get("model") or "unknown").strip()
            op    = (rec.get("operation") or "unknown").strip()
            if inp or out:
                pk = resolve_pricing_key(model)
                pricing = get_effective_model_pricing().get(pk, {}) if pk else {}
                cost = (inp * pricing.get("input", 0) + out * pricing.get("output", 0)) / 1_000_000
            elif secs:
                cost = get_audio_cost(model, secs)
            else:
                cost = 0.0
            try:
                rec_date = datetime.fromisoformat(
                    (rec.get("datetime") or "").replace("Z", "+00:00")
                ).date()
            except Exception:
                rec_date = today_d
            delta = (today_d - rec_date).days

            for b in [totals] + ([today_b] if delta == 0 else []) \
                               + ([last7]   if delta < 7  else []) \
                               + ([last30]  if delta < 30 else []):
                b["input"] += inp; b["output"] += out; b["cost"] += cost

            for d, key in [(by_day, rec_date.isoformat()), (by_model, model), (by_op, op)]:
                if key not in d:
                    d[key] = _eb()
                d[key]["input"] += inp; d[key]["output"] += out; d[key]["cost"] += cost

        active_days = max(len(by_day), 1)
        return {
            "totals":   totals,
            "today":    today_b,
            "last7":    last7,
            "last30":   last30,
            "daily_avg": {k: totals[k] / active_days for k in totals},
            "by_day":    dict(sorted(by_day.items())),
            "by_model":  dict(sorted(by_model.items(),  key=lambda x: -x[1]["cost"])),
            "by_op":     dict(sorted(by_op.items(),     key=lambda x: -x[1]["cost"])),
        }

    conn = sqlite3.connect(DB_FILE)

    cur = conn.cursor()
    cur.execute("SELECT email, first_name, last_name FROM users WHERE email = ?", (session_email,))
    _u = cur.fetchone()
    current_user = {"email": _u[0], "first_name": _u[1], "last_name": _u[2]} if _u else {"email": session_email, "first_name": "", "last_name": ""}

    # ── per-user leaderboard (admin only) ─────────────────────────────────────
    user_leaderboard = []
    all_users = []
    if is_admin:
        c = conn.cursor()
        c.execute("SELECT email, first_name, last_name FROM users ORDER BY email")
        all_users = [{"email": r[0], "name": f"{r[1]} {r[2]}".strip() or r[0]} for r in c.fetchall()]

        for u in all_users:
            recs = _fetch_records(conn, u["email"])
            agg  = _aggregate(recs)
            user_leaderboard.append({
                "email":  u["email"],
                "name":   u["name"],
                "input":  agg["totals"]["input"],
                "output": agg["totals"]["output"],
                "cost":   agg["totals"]["cost"],
            })
        user_leaderboard.sort(key=lambda x: -x["cost"])

    # ── selected view ──────────────────────────────────────────────────────────
    email_filter = None if (is_admin and view_email == "__all__") else view_email
    records = _fetch_records(conn, email_filter)
    conn.close()

    agg = _aggregate(records)
    view_label = "Wszyscy użytkownicy" if view_email == "__all__" else view_email

    return render_template(
        'usage-history-page.html',
        **agg,
        model_pricing=get_effective_model_pricing(),
        is_admin=is_admin,
        view_email=view_email,
        view_label=view_label,
        all_users=all_users,
        user_leaderboard=user_leaderboard,
        user=current_user,
    )

@app.route('/settings', methods=['GET'])
def settings():
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))

    api_key_status = {
        provider: {
            "label": config["label"],
            "env_key": config["env_key"],
            "configured": bool(get_provider_api_key(provider)),
            "value": get_provider_api_key(provider) if provider == "ollama_ip" else ""
        }
        for provider, config in PROVIDERS.items()
    }

    _, pricing_last_updated = load_pricing_data()
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("SELECT email, first_name, last_name FROM users WHERE email = ?", (session['user_email'],))
    _u = cur.fetchone()
    conn.close()
    current_user = {"email": _u[0], "first_name": _u[1], "last_name": _u[2]} if _u else {"email": session['user_email'], "first_name": "", "last_name": ""}
    return render_template(
        'settings.html',
        models=list_ai_models(),
        providers=PROVIDERS,
        model_types=MODEL_TYPES,
        model_catalog=MODEL_CATALOG,
        local_model_requirements=LOCAL_MODEL_REQUIREMENTS,
        api_key_status=api_key_status,
        model_pricing=get_effective_model_pricing(),
        audio_pricing=AUDIO_PRICING,
        pricing_last_updated=pricing_last_updated,
        is_admin=(session.get('user_email') == ADMIN_EMAIL),
        user=current_user,
    )

@app.route('/admin/restart', methods=['POST'])
def admin_restart():
    if 'user_email' not in session:
        return jsonify({"error": "Brak autoryzacji"}), 401
    if session['user_email'] != ADMIN_EMAIL:
        return jsonify({"error": "Brak uprawnień administratora"}), 403

    try:
        app.logger.warning("[admin/restart] Restart zainicjowany przez %s", session['user_email'])
        restart_server()
        return jsonify({"ok": True, "message": "Serwer restartuje się. Odśwież stronę za kilka sekund."})
    except Exception as exc:
        app.logger.error("[admin/restart] Nieoczekiwany błąd: %s", exc)
        return jsonify({"error": f"Błąd: {exc}"}), 500

@app.route('/settings/update-pricing', methods=['POST'])
def update_pricing_route():
    if 'user_email' not in session or session['user_email'] != ADMIN_EMAIL:
        return jsonify({'ok': False, 'error': 'Brak uprawnień'}), 403

    LITELLM_URL = 'https://raw.githubusercontent.com/BerriAI/litellm/main/model_prices_and_context_window.json'
    try:
        resp = requests.get(LITELLM_URL, timeout=15)
        resp.raise_for_status()
        litellm_data = resp.json()
    except Exception as exc:
        return jsonify({'ok': False, 'error': f'Nie udało się pobrać cennika: {exc}'}), 502

    current_overrides, _ = load_pricing_data()
    effective = get_effective_model_pricing()

    changes = []
    new_overrides = dict(current_overrides)

    for model_key in MODEL_PRICING:
        # LiteLLM may use bare name, openai/ prefix, or groq/ prefix
        litellm_entry = None
        for candidate in [model_key, f'openai/{model_key}', f'groq/{model_key}']:
            if candidate in litellm_data:
                litellm_entry = litellm_data[candidate]
                break
        if not litellm_entry:
            continue

        inp_per_tok = litellm_entry.get('input_cost_per_token')
        out_per_tok = litellm_entry.get('output_cost_per_token')
        if inp_per_tok is None or out_per_tok is None:
            continue

        new_input  = round(float(inp_per_tok) * 1_000_000, 4)
        new_output = round(float(out_per_tok) * 1_000_000, 4)
        old_input  = effective.get(model_key, MODEL_PRICING[model_key])['input']
        old_output = effective.get(model_key, MODEL_PRICING[model_key])['output']

        if abs(new_input - old_input) > 0.001 or abs(new_output - old_output) > 0.001:
            changes.append({
                'model': model_key,
                'old_input': old_input,  'new_input': new_input,
                'old_output': old_output, 'new_output': new_output,
            })
            new_overrides[model_key] = {'input': new_input, 'output': new_output}

    timestamp = None
    if changes:
        timestamp = datetime.now(timezone.utc).timestamp()
        save_pricing_data(new_overrides, timestamp)

    return jsonify({'ok': True, 'changes': changes, 'timestamp': timestamp})

@app.route('/settings/models', methods=['POST'])
def add_model():
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))

    provider = request.form.get('provider', '').strip().lower()
    model_type = request.form.get('model_type', '').strip().lower()
    display_name = request.form.get('display_name', '').strip()
    model_id = request.form.get('model_id', '').strip()
    enabled = 1 if request.form.get('enabled') == 'on' else 0
    is_default = 1 if request.form.get('is_default') == 'on' else 0

    if not enabled:
        is_default = 0

    if provider not in PROVIDERS or model_type not in MODEL_TYPES or not display_name or not model_id:
        flash('Uzupełnij poprawnie providera, typ, nazwę oraz identyfikator modelu.', 'danger')
        return redirect(url_for('settings'))

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()

    if is_default:
        cursor.execute(
            "UPDATE ai_models SET is_default = 0 WHERE provider = ? AND model_type = ?",
            (provider, model_type)
        )

    cursor.execute(
        """
        INSERT INTO ai_models (provider, model_type, display_name, model_id, enabled, is_default)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (provider, model_type, display_name, model_id, enabled, is_default)
    )
    conn.commit()
    conn.close()

    flash('Model został dodany.', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/models/<int:model_pk>', methods=['POST'])
def update_model(model_pk):
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))

    provider = request.form.get('provider', '').strip().lower()
    model_type = request.form.get('model_type', '').strip().lower()
    display_name = request.form.get('display_name', '').strip()
    model_id = request.form.get('model_id', '').strip()
    enabled = 1 if request.form.get('enabled') == 'on' else 0
    is_default = 1 if request.form.get('is_default') == 'on' else 0

    if display_name.lower() == "usun model":
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT provider, model_type, display_name, is_default FROM ai_models WHERE id = ?", (model_pk,))
        model_row = cursor.fetchone()
        if not model_row:
            conn.close()
            flash('Nie znaleziono modelu do usunięcia.', 'danger')
            return redirect(url_for('settings'))

        deleted_provider, deleted_model_type, deleted_display_name, was_default = model_row
        cursor.execute("DELETE FROM ai_models WHERE id = ?", (model_pk,))
        if was_default:
            promote_first_enabled_model_as_default(cursor, deleted_provider, deleted_model_type)

        conn.commit()
        conn.close()
        flash(f"Model '{deleted_display_name}' został usunięty.", 'success')
        return redirect(url_for('settings'))

    if not enabled:
        is_default = 0

    if provider not in PROVIDERS or model_type not in MODEL_TYPES or not display_name or not model_id:
        flash('Nie zapisano zmian. Sprawdź dane modelu.', 'danger')
        return redirect(url_for('settings'))

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()

    if is_default:
        cursor.execute(
            "UPDATE ai_models SET is_default = 0 WHERE provider = ? AND model_type = ? AND id != ?",
            (provider, model_type, model_pk)
        )

    cursor.execute(
        """
        UPDATE ai_models
        SET provider = ?,
            model_type = ?,
            display_name = ?,
            model_id = ?,
            enabled = ?,
            is_default = ?,
            updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """,
        (provider, model_type, display_name, model_id, enabled, is_default, model_pk)
    )
    conn.commit()
    conn.close()

    flash('Model został zaktualizowany.', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/ollama-ip-url', methods=['POST'])
def update_ollama_ip_url():
    if 'user_email' not in session:
        flash('Brak dostępu. Musisz się najpierw zalogować!', 'danger')
        return redirect(url_for('login'))
    url = request.form.get('ollama_ip_url', '').strip().rstrip('/')
    if url and not url.startswith(('http://', 'https://')):
        flash('Nieprawidłowy URL — musi zaczynać się od http:// lub https://', 'danger')
        return redirect(url_for('settings'))
    write_env_var('OLLAMA_IP_URL', url)
    if url:
        flash(f'URL Ollama zapisany: {url}', 'success')
    else:
        flash('URL Ollama wyczyszczony.', 'success')
    return redirect(url_for('settings'))

@app.route('/api/ollama-ip/status', methods=['GET'])
@limiter.limit("30 per minute")
def ollama_ip_status():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Unauthorized"}), 401

    base_url = get_provider_api_key("ollama_ip").rstrip("/")
    if not base_url:
        return jsonify({"available": False, "reason": "no_url", "url": None})

    try:
        resp = requests.get(f"{base_url}/api/tags", timeout=3)
        resp.raise_for_status()
        data = resp.json()
        models_list = [m.get("name") for m in data.get("models", [])]
        return jsonify({"available": True, "url": base_url, "models": models_list})
    except requests.exceptions.ConnectionError:
        return jsonify({"available": False, "reason": "connection_error", "url": base_url})
    except requests.exceptions.Timeout:
        return jsonify({"available": False, "reason": "timeout", "url": base_url})
    except Exception:
        return jsonify({"available": False, "reason": "error", "url": base_url})


@app.route('/api/models', methods=['GET'])
def api_models():
    print
    model_type = (request.args.get("type") or request.args.get("model_type") or "").strip().lower() or None
    if model_type and model_type not in MODEL_TYPES:
        allowed_types = "', '".join(MODEL_TYPES.keys())
        return jsonify({"error": f"Nieobsługiwany typ modelu. Użyj jednego z: '{allowed_types}'."}), 400

    include_disabled = parse_bool_setting(request.args.get("include_disabled"), default=False)

    local_models = build_local_model_list(model_type=model_type)
    cloud_models = build_cloud_model_list(model_type=model_type, include_disabled=include_disabled)

    return jsonify({
        "local": local_models,
        "cloud": cloud_models,
        "models": local_models + cloud_models,
        "filters": {
            "type": model_type,
            "include_disabled": include_disabled
        },
        "counts": {
            "local": len(local_models),
            "cloud": len(cloud_models),
            "total": len(local_models) + len(cloud_models)
        }
    })

@app.route('/get-history', methods=['GET'])
def get_history():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    try:
        offset = max(0, int(request.args.get('offset', 0)))
    except (ValueError, TypeError):
        offset = 0

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()

    cursor.execute(
        "SELECT COUNT(*) FROM history WHERE user_email = ?",
        (user_email,)
    )
    total = cursor.fetchone()[0]

    cursor.execute(
        """
        SELECT id, filename, raw_text, ai_notes, COALESCE(notes_model_used, ''), COALESCE(openai_usage_history, ''), datetime(created_at, 'localtime'), project_id
        FROM history
        WHERE user_email = ?
        ORDER BY created_at DESC
        LIMIT ? OFFSET ?
        """,
        (user_email, HISTORY_PAGE_SIZE, offset)
    )
    rows = cursor.fetchall()

    record_ids = [r[0] for r in rows]
    chat_rows = []
    if record_ids:
        placeholders = ','.join('?' * len(record_ids))
        cursor.execute(
            f"""
            SELECT record_id, role, content, datetime(created_at, 'localtime'), COALESCE(model_used, ''), COALESCE(openai_usage_history, '')
            FROM chat_history
            WHERE record_id IN ({placeholders})
            ORDER BY id ASC
            """,
            record_ids
        )
        chat_rows = cursor.fetchall()
    conn.close()

    chat_by_record = {}
    for record_id, role, content, created_at, model_used, openai_usage_history in chat_rows:
        chat_by_record.setdefault(record_id, []).append({
            "role": role,
            "content": content,
            "created_at": created_at,
            "engine": model_used,
            "openai_usage_history": parse_openai_usage_history(openai_usage_history)
        })
    
    history_list = []
    for r in rows:
        chat_messages = chat_by_record.get(r[0], [])
        notes_usage = parse_openai_usage_history(r[5])
        chat_usage_lists = [m.get("openai_usage_history", []) for m in chat_messages]
        history_list.append({
            "id": r[0],
            "filename": r[1],
            "raw_text": r[2],
            "ai_notes": r[3],
            "notes_model_used": r[4],
            "openai_usage_history": notes_usage,
            "created_at": r[6],
            "chat_messages": chat_messages,
            "chat_count": len(chat_messages),
            "project_id": r[7],
            "authenticity_score": extract_authenticity_score(r[3]),
            "token_summary": compute_token_summary([notes_usage] + chat_usage_lists),
        })
    return jsonify({"items": history_list, "total": total, "offset": offset, "limit": HISTORY_PAGE_SIZE})


# Trasy dla poszczególnych API endpoints uwzględniające ograniczenia limitów i autoryzację
@app.route('/transcribe', methods=['POST'])
@limiter.limit("30 per hour; 5 per minute")
@csrf.exempt
def transcribe():
    # print(f"[API /transcribe] Request received with form data: {request.form} and files: {request.files}")
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
        
    youtube_url = request.form.get('youtube_url', '').strip()
    webpage_url = request.form.get('webpage_url', '').strip()
    direct_text = request.form.get('direct_text', '').strip()
    processing_mode = request.form.get('processing_mode', get_default_processing_mode())
    model_name = request.form.get('model_name', 'base')
    cloud_model_id = request.form.get('cloud_model_id')
    language = request.form.get('language', 'auto')
    task = request.form.get('task', 'transcribe')
    custom_name = request.form.get('custom_name', '').strip()
    notes_mode = request.form.get('notes_mode', 'full').strip()
    if notes_mode not in {'full', 'summary', 'overview', 'bullets', 'prompt', 'reel-prepare'}:
        notes_mode = 'full'
    notes_model_id = request.form.get('notes_model_id') or None
    project_id_form = None
    _raw_pid = request.form.get('project_id')
    if _raw_pid:
        try:
            project_id_form = int(_raw_pid)
        except (ValueError, TypeError):
            project_id_form = None

    file_path = None
    display_title = youtube_url or webpage_url or custom_name or "Przesłany plik"
    notes_model_used = ""
    yt_transcript_text = None

    try:
        if direct_text:
            if len(direct_text) < 3:
                return jsonify({"error": "Tekst jest zbyt krótki"}), 400
            if _is_multi_url_text(direct_text):
                combined, auto_title = fetch_multi_url_content(direct_text)
                direct_text = combined
                display_title = custom_name if custom_name else auto_title
            else:
                display_title = custom_name if custom_name else "Tekst wklejony"
        elif youtube_url:
            try:
                app.logger.info("[transcribe] START ▶️YouTube 📥 try fetch_youtube_transcript dla %s", youtube_url)
                yt_api_result = fetch_youtube_transcript(youtube_url)
                yt_transcript_text = yt_api_result["text"]
                display_title = custom_name if custom_name else yt_api_result["title"]
                app.logger.info("[transcribe] YouTube Transcript API: pobrano napisy dla %s", yt_api_result["video_id"])
            except Exception as yt_api_err:
                app.logger.info("[transcribe] YouTube Transcript API niedostępne (%s)", yt_api_err)
                try:
                    app.logger.info("[transcribe] START ▶️YouTube 📥 try download_youtube_audio dla %s", youtube_url)
                    youtube_download = download_youtube_audio(youtube_url)
                    file_path = youtube_download["file_path"]
                    display_title = custom_name if custom_name else f"YT: {youtube_download['title']}"
                except Exception as yt_download_err:
                    app.logger.error("[transcribe] Błąd pobierania audio z YouTube (%s)", yt_download_err)
                    # raise Exception("Nie można pobrać transkrypcji ani audio z podanego linku YouTube.")    
                
        elif webpage_url:
            _web_content, _web_title = fetch_webpage_content(webpage_url)
            display_title = custom_name if custom_name else (_web_title or webpage_url)
        else:
            if 'file' not in request.files:
                return jsonify({"error": "Brak pliku, linku YouTube lub adresu URL strony"}), 400
            file = request.files['file']
            if file.filename == '':
                return jsonify({"error": "Nie wybrano pliku"}), 400
            ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
            if ext not in ALLOWED_AUDIO_EXTENSIONS:
                return jsonify({"error": f"Niedozwolony typ pliku '.{ext}'. Akceptowane: {', '.join(sorted(ALLOWED_AUDIO_EXTENSIONS))}"}), 400

            file_path = make_temp_upload_path(file.filename)
            file.save(file_path)
            display_title = custom_name if custom_name else file.filename
    except Exception as e:
        print(f"Error during file handling or content fetching: {e}")
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({"error": str(e)}), 400

    try:
        # Obsługa transkrypcji strony internetowej lub z pliku tekstowego bez STT
        if direct_text:
            surowy_tekst = direct_text
            detected_lang = "Tekst wklejony"
            model_used_info = "Tekst wklejony (brak STT)"
            notes_model_used = describe_notes_model(None, processing_mode)
            if processing_mode == 'online':
                selected_transcription_model = get_selected_transcription_model(cloud_model_id)
                preferred_provider = selected_transcription_model["provider"] if selected_transcription_model else None
                notatki_ai, notes_model = generate_audio_notes(
                    surowy_tekst, processing_mode,
                    preferred_provider=preferred_provider, model_used=None,
                    notes_mode=notes_mode, notes_model_id=notes_model_id
                )
                notes_model_used = describe_notes_model(notes_model, processing_mode)
            else:
                try:
                    prompt = build_audio_notes_prompt(surowy_tekst, mode=notes_mode)
                    response = ollama.chat(model='llama3', messages=[{'role': 'user', 'content': prompt}])
                    notatki_ai = response['message']['content']
                except Exception:
                    notatki_ai = ""
        elif yt_transcript_text:
            surowy_tekst = yt_transcript_text
            detected_lang = "YouTube napisy"
            model_used_info = "YouTube Transcript API (napisy)"
            notes_model_used = describe_notes_model(None, processing_mode)
            if processing_mode == 'online':
                selected_transcription_model = get_selected_transcription_model(cloud_model_id)
                preferred_provider = selected_transcription_model["provider"] if selected_transcription_model else None
                notatki_ai, notes_model = generate_audio_notes(
                    surowy_tekst, processing_mode,
                    preferred_provider=preferred_provider, model_used=None,
                    notes_mode=notes_mode, notes_model_id=notes_model_id
                )
                notes_model_used = describe_notes_model(notes_model, processing_mode)
            else:
                try:
                    prompt = build_audio_notes_prompt(surowy_tekst, mode=notes_mode)
                    response = ollama.chat(model='llama3', messages=[{'role': 'user', 'content': prompt}])
                    notatki_ai = response['message']['content']
                except Exception:
                    notatki_ai = ""
        elif webpage_url or (file_path and file_path.endswith('.txt')):
            if webpage_url:
                surowy_tekst = _web_content
                detected_lang = "Strona internetowa"
                model_used_info = "Pobranie treści strony"
                notes_model_used = describe_notes_model(None, processing_mode)

            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    surowy_tekst = f.read()
                    if os.path.exists(file_path):
                        os.remove(file_path)

                detected_lang = "Plik tekstowy"
                model_used_info = "Czysty tekst (Brak STT)"
                notes_model_used = describe_notes_model(None, processing_mode)

            if processing_mode == 'online':
                selected_transcription_model = get_selected_transcription_model(cloud_model_id)
                preferred_provider = selected_transcription_model["provider"] if selected_transcription_model else None
                notatki_ai, notes_model = generate_audio_notes(
                    surowy_tekst, processing_mode,
                    preferred_provider=preferred_provider, model_used=None,
                    notes_mode=notes_mode, notes_model_id=notes_model_id
                )
                notes_model_used = describe_notes_model(notes_model, processing_mode)
                model_used_info = f"Strona internetowa + notatki: {describe_cloud_model(notes_model)}"
            else:
                try:
                    prompt = build_audio_notes_prompt(surowy_tekst, mode=notes_mode)
                    response = ollama.chat(model='llama3', messages=[{'role': 'user', 'content': prompt}])
                    notatki_ai = response['message']['content']
                except Exception:
                    notatki_ai = "Nie udało się wygenerować notatek AI lokalnie. Upewnij się, że Ollama działa w tle."

        else:
            # Klasyczne przetwarzanie audio STT (plik lub YouTube)
            transcription_result = process_audio_transcription(
                file_path,
                processing_mode=processing_mode,
                model_name=model_name,
                cloud_model_id=cloud_model_id,
                language=language,
                task=task,
                notes_mode=notes_mode,
                notes_model_id=notes_model_id
            )
            surowy_tekst = transcription_result["text"]
            notatki_ai = transcription_result["notes"]
            notes_model_used = transcription_result.get("notes_model_used", "")
            detected_lang = transcription_result["language"]
            task = transcription_result["task"]
            model_used_info = transcription_result["model_used"]

            if os.path.exists(file_path):
                os.remove(file_path)

        try:
            openai_usage_history = get_current_openai_usage_history()
            new_id, new_project_id = save_transcription_history(
                user_email,
                display_title,
                surowy_tekst,
                notatki_ai,
                notes_model_used,
                openai_usage_history,
                project_id=project_id_form
            )
        except Exception as e:
            print(f"Nie udało się zapisać historii transkrypcji: {e}")
            new_id = None
            new_project_id = None

        return jsonify({
            "text": surowy_tekst,
            "notes": notatki_ai,
            "notes_model_used": notes_model_used,
            "model_used": model_used_info,
            "language": detected_lang,
            "task": task,
            "saved_name": display_title,
            "record_id": new_id,
            "project_id": new_project_id,
            "openai_usage_history": openai_usage_history,
            "authenticity_score": extract_authenticity_score(notatki_ai),
        })

    except ValueError as e:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        log_app_error(user_email, 'transcription', str(e), level='warning')
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        import traceback
        log_app_error(user_email, 'transcription', str(e), traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route('/api/youtube/transcribe', methods=['POST'])
@limiter.limit("10 per hour; 2 per minute")
@csrf.exempt
def api_youtube_transcribe():
    # print(f"[API /api/youtube/transcribe] Request received with JSON: {request.get_json(silent=True)}")
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    payload = request.get_json(silent=True)
    if not isinstance(payload, dict):
        return jsonify({"error": "Oczekiwano danych JSON"}), 400

    settings = payload.get("settings") or {}
    if not isinstance(settings, dict):
        return jsonify({"error": "Pole settings musi być obiektem JSON"}), 400

    yt_url_value = (
        payload.get("yt_url")
        or payload.get("youtube_url")
        or settings.get("yt_url")
        or settings.get("youtube_url")
        or ""
    )
    yt_url = str(yt_url_value).strip()
    if not yt_url:
        return jsonify({"error": "Brak wymaganego pola yt_url"}), 400

    processing_mode = get_payload_setting(payload, settings, "processing_mode", get_default_processing_mode())
    model_name = get_payload_setting(payload, settings, "model_name", "base")
    cloud_model_id = get_payload_setting(payload, settings, "cloud_model_id")
    language = get_payload_setting(payload, settings, "language", "auto")
    task = get_payload_setting(payload, settings, "task", "transcribe")
    custom_name = str(get_payload_setting(payload, settings, "custom_name", "") or "").strip()
    save_to_history = parse_bool_setting(get_payload_setting(payload, settings, "save_to_history"), default=False)

    file_path = None

    try:
        # Próba szybkiego pobrania napisów (bez pobierania audio)
        yt_transcript = None
        youtube_download = None
        try:
            yt_transcript = fetch_youtube_transcript(yt_url)
            app.logger.info("YouTube Transcript API: pobrano napisy dla %s", yt_transcript["video_id"])
        except Exception as yt_api_err:
            app.logger.info("YouTube Transcript API niedostępne (%s), pobieranie audio...", yt_api_err)
            youtube_download = download_youtube_audio(yt_url)
            file_path = youtube_download["file_path"]

        if yt_transcript:
            # Ścieżka przez Transcript API — brak STT, tylko notatki
            selected_transcription_model = get_selected_transcription_model(cloud_model_id)
            preferred_provider = selected_transcription_model["provider"] if selected_transcription_model else None
            notatki_ai, notes_model = generate_audio_notes(
                yt_transcript["text"], processing_mode,
                preferred_provider=preferred_provider, model_used=None
            )
            notes_model_used = describe_notes_model(notes_model, processing_mode)
            saved_name = custom_name if custom_name else yt_transcript["title"]
            record_id = None
            project_id = None
            openai_usage_history = get_current_openai_usage_history()
            if save_to_history:
                record_id, project_id = save_transcription_history(
                    user_email, saved_name,
                    yt_transcript["text"], notatki_ai, notes_model_used, openai_usage_history
                )
            return jsonify({
                "text": yt_transcript["text"],
                "summary": notatki_ai,
                "notes": notatki_ai,
                "language": "YouTube napisy",
                "model_used": "YouTube Transcript API (napisy)",
                "notes_model_used": notes_model_used,
                "task": "transcript",
                "saved": record_id is not None,
                "saved_name": saved_name,
                "record_id": record_id,
                "project_id": project_id,
                "openai_usage_history": openai_usage_history,
                "authenticity_score": extract_authenticity_score(notatki_ai),
                "youtube": {"id": yt_transcript["video_id"], "title": yt_transcript["title"],
                            "duration": None, "url": yt_url}
            })

        # Fallback: audio pobrany przez yt_dlp → STT
        transcription_result = process_audio_transcription(
            file_path,
            processing_mode=processing_mode,
            model_name=model_name,
            cloud_model_id=cloud_model_id,
            language=language,
            task=task
        )

        if os.path.exists(file_path):
            os.remove(file_path)

        saved_name = custom_name if custom_name else f"YT: {youtube_download['title']}"
        record_id = None
        project_id = None
        openai_usage_history = get_current_openai_usage_history()
        if save_to_history:
            record_id, project_id = save_transcription_history(
                user_email,
                saved_name,
                transcription_result["text"],
                transcription_result["notes"],
                transcription_result.get("notes_model_used", ""),
                openai_usage_history
            )

        return jsonify({
            "text": transcription_result["text"],
            "summary": transcription_result["notes"],
            "notes": transcription_result["notes"],
            "language": transcription_result["language"],
            "model_used": transcription_result["model_used"],
            "notes_model_used": transcription_result.get("notes_model_used", ""),
            "task": transcription_result["task"],
            "saved": record_id is not None,
            "saved_name": saved_name,
            "record_id": record_id,
            "project_id": project_id,
            "openai_usage_history": openai_usage_history,
            "authenticity_score": extract_authenticity_score(transcription_result["notes"]),
            "youtube": {
                "id": youtube_download["id"],
                "title": youtube_download["title"],
                "duration": youtube_download["duration"],
                "url": youtube_download["webpage_url"]
            }
        })

    except ValueError as e:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        log_app_error(user_email, 'youtube', str(e), level='warning')
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        import traceback as _tb
        log_app_error(user_email, 'youtube', str(e), _tb.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route('/api/webpage/read', methods=['POST'])
@csrf.exempt
def api_webpage_read():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    payload = request.get_json(silent=True)
    if not isinstance(payload, dict):
        return jsonify({"error": "Oczekiwano danych JSON"}), 400

    settings = payload.get("settings") or {}
    if not isinstance(settings, dict):
        return jsonify({"error": "Pole settings musi być obiektem JSON"}), 400

    url_value = (
        payload.get("url")
        or payload.get("webpage_url")
        or settings.get("url")
        or settings.get("webpage_url")
        or ""
    )
    url = str(url_value).strip()
    if not url:
        return jsonify({"error": "Brak wymaganego pola url"}), 400

    processing_mode = get_payload_setting(payload, settings, "processing_mode", get_default_processing_mode())
    cloud_model_id = get_payload_setting(payload, settings, "cloud_model_id")
    custom_name = str(get_payload_setting(payload, settings, "custom_name", "") or "").strip()
    save_to_history = parse_bool_setting(get_payload_setting(payload, settings, "save_to_history"), default=False)

    try:
        web_text, web_title = fetch_webpage_content(url)
        saved_name = custom_name if custom_name else (web_title or url)
        notes_model_used = describe_notes_model(None, processing_mode)

        if processing_mode == 'online':
            selected_transcription_model = get_selected_transcription_model(cloud_model_id)
            preferred_provider = selected_transcription_model["provider"] if selected_transcription_model else None
            notes, notes_model = generate_audio_notes(
                web_text, 
                processing_mode, 
                preferred_provider=preferred_provider,
                model_used=None
            )
            notes_model_used = describe_notes_model(notes_model, processing_mode)
            model_used_info = f"Strona internetowa + notatki: {describe_cloud_model(notes_model)}"
        else:
            try:
                prompt = build_audio_notes_prompt(web_text)
                local_response = ollama.chat(model='llama3', messages=[{'role': 'user', 'content': prompt}])
                notes = local_response['message']['content']
            except:
                notes = "Nie udało się wygenerować notatek AI lokalnie. Upewnij się, że Ollama działa w tle."
            model_used_info = "Strona internetowa + Ollama (Llama 3)"

        record_id = None
        project_id = None
        openai_usage_history = get_current_openai_usage_history()
        if save_to_history:
            record_id, project_id = save_transcription_history(
                user_email,
                saved_name,
                web_text,
                notes,
                notes_model_used,
                openai_usage_history
            )

        return jsonify({
            "text": web_text,
            "summary": notes,
            "notes": notes,
            "notes_model_used": notes_model_used,
            "model_used": model_used_info,
            "saved": record_id is not None,
            "saved_name": saved_name,
            "record_id": record_id,
            "project_id": project_id,
            "openai_usage_history": openai_usage_history,
            "authenticity_score": extract_authenticity_score(notes),
            "webpage": {
                "url": url,
                "title": web_title
            }
        })

    except ValueError as e:
        log_app_error(user_email, 'ask-question', str(e), level='warning')
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        import traceback as _tb
        log_app_error(user_email, 'ask-question', str(e), _tb.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route('/new-chat', methods=['POST'])
@limiter.limit("30 per hour")
@csrf.exempt
def new_chat():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    record_id, project_id = save_transcription_history(user_email, 'Nowy czat', '', '')
    return jsonify({"record_id": record_id, "project_id": project_id})

@app.route('/ask-question', methods=['POST'])
@limiter.limit("60 per hour; 10 per minute")
@csrf.exempt
def ask_question():
    # print(f"[API /ask-question] Request received with JSON: {request.get_json(silent=True)}")
    user_email = get_authenticated_user()
    if not user_email:
        app.logger.warning("ask-question: unauthorized request from %s", request.remote_addr)
        return jsonify({"error": "Brak autoryzacji"}), 401
        
    data = request.get_json(silent=True) or {}
    if not isinstance(data, dict):
        app.logger.warning(
            "ask-question: invalid json body user=%s body_type=%s",
            user_email,
            type(data).__name__
        )
        return jsonify({"error": "Nieprawidłowy JSON"}), 400

    record_id = data.get('id')
    project_id = data.get('project_id')
    question = str(data.get('question') or '').strip()

    if not (record_id or project_id) or not question:
        app.logger.warning("ask-question: invalid payload user=%s", user_email)
        return jsonify({"error": "Brak ID nagrania/projektu lub pytania"}), 400

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    use_project_chat = False

    if project_id:
        # project path — verify ownership
        cursor.execute("SELECT id FROM projects WHERE id = ? AND user_email = ?", (project_id, user_email))
        if not cursor.fetchone():
            conn.close()
            return jsonify({"error": "Nie znaleziono projektu"}), 404
        sources = get_project_sources(project_id, user_email)
        if not sources:
            conn.close()
            return jsonify({"error": "Projekt nie ma żadnych źródeł"}), 404
        if len(sources) == 1:
            record_id = sources[0]['id']
            transkrypcja = sources[0]['raw_text']
        else:
            transkrypcja = build_project_context(sources)
            use_project_chat = True
        if use_project_chat:
            cursor.execute(
                "SELECT role, content FROM (SELECT role, content, id FROM project_chat_history WHERE project_id = ? ORDER BY id DESC LIMIT 6) ORDER BY id ASC",
                (project_id,)
            )
        else:
            cursor.execute(
                "SELECT role, content FROM (SELECT role, content, id FROM chat_history WHERE record_id = ? ORDER BY id DESC LIMIT 6) ORDER BY id ASC",
                (record_id,)
            )
    else:
        # legacy single-record path
        cursor.execute("SELECT raw_text FROM history WHERE id = ? AND user_email = ?", (record_id, user_email))
        row = cursor.fetchone()
        if not row:
            conn.close()
            app.logger.warning("ask-question: record not found user=%s record_id=%s", user_email, record_id)
            return jsonify({"error": "Nie znaleziono nagrania w Twojej historii"}), 404
        transkrypcja = row[0]
        cursor.execute(
            "SELECT role, content FROM (SELECT role, content, id FROM chat_history WHERE record_id = ? ORDER BY id DESC LIMIT 6) ORDER BY id ASC",
            (record_id,)
        )

    context_rows = cursor.fetchall()
    conn.close()

    try:
        response_prompt = build_web_search_question_prompt(transkrypcja, context_rows, question)
        chat_model_id = data.get('chat_model_id')
        chat_model = get_ai_model_by_id(chat_model_id, "chat") if chat_model_id else None
        if not chat_model:
            chat_model = get_default_ai_model("chat")
        ai_answer = answer_question_with_ai(response_prompt, chat_model)
        odpowiedz_ai = ai_answer["text"]
        openai_usage_history = get_current_openai_usage_history()

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        if use_project_chat:
            cursor.execute("INSERT INTO project_chat_history (project_id, role, content) VALUES (?, ?, ?)", (project_id, 'user', question))
            cursor.execute(
                "INSERT INTO project_chat_history (project_id, role, content, model_used) VALUES (?, ?, ?, ?)",
                (project_id, 'assistant', odpowiedz_ai, ai_answer["engine"])
            )
        else:
            cursor.execute("INSERT INTO chat_history (record_id, role, content) VALUES (?, ?, ?)", (record_id, 'user', question))
            cursor.execute(
                "INSERT INTO chat_history (record_id, role, content, model_used, openai_usage_history) VALUES (?, ?, ?, ?, ?)",
                (record_id, 'assistant', odpowiedz_ai, ai_answer["engine"], serialize_openai_usage_history(openai_usage_history))
            )
        conn.commit()
        conn.close()

        return jsonify({
            "answer": odpowiedz_ai,
            "engine": ai_answer["engine"],
            "chat_model_used": ai_answer["engine"],
            "sources": ai_answer["sources"],
            "openai_usage_history": openai_usage_history,
            "authenticity_score": extract_authenticity_score(odpowiedz_ai),
        })
    except Exception as e:
        import traceback as _tb
        app.logger.exception("ask-question: error user=%s project_id=%s record_id=%s", user_email, project_id, record_id)
        log_app_error(user_email, 'ai-chat', f"Błąd AI: {str(e)}", _tb.format_exc())
        return jsonify({"error": f"Błąd AI: {str(e)}"}), 500

@app.route('/delete-history/<int:item_id>', methods=['DELETE'])
@csrf.exempt
def delete_history(item_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT project_id FROM history WHERE id = ? AND user_email = ?", (item_id, user_email))
    row = cursor.fetchone()
    if row:
        pid = row[0]
        cursor.execute("DELETE FROM history WHERE id = ? AND user_email = ?", (item_id, user_email))
        if pid:
            cursor.execute("SELECT COUNT(*) FROM history WHERE project_id = ?", (pid,))
            if cursor.fetchone()[0] == 0:
                cursor.execute("DELETE FROM projects WHERE id = ? AND user_email = ?", (pid, user_email))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route('/delete-history/bulk', methods=['POST'])
@csrf.exempt
def delete_history_bulk():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    ids = data.get('ids', [])
    if not isinstance(ids, list) or not ids:
        return jsonify({"error": "Brak listy ID"}), 400
    ids = [int(i) for i in ids if str(i).isdigit()]
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    placeholders = ','.join('?' * len(ids))
    cursor.execute(
        f"SELECT DISTINCT project_id FROM history WHERE id IN ({placeholders}) AND user_email = ? AND project_id IS NOT NULL",
        ids + [user_email]
    )
    affected_projects = [r[0] for r in cursor.fetchall()]
    cursor.executemany(
        "DELETE FROM history WHERE id = ? AND user_email = ?",
        [(i, user_email) for i in ids]
    )
    for pid in affected_projects:
        cursor.execute("SELECT COUNT(*) FROM history WHERE project_id = ?", (pid,))
        if cursor.fetchone()[0] == 0:
            cursor.execute("DELETE FROM projects WHERE id = ? AND user_email = ?", (pid, user_email))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "deleted": len(ids)})

@app.route('/history/<int:item_id>/rename', methods=['PATCH'])
@csrf.exempt
def rename_history(item_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    new_name = str(data.get('name') or '').strip()
    if not new_name:
        return jsonify({"error": "Pusty tytuł"}), 400
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE history SET filename = ? WHERE id = ? AND user_email = ?",
        (new_name, item_id, user_email)
    )
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route('/export/docx', methods=['POST'])
@csrf.exempt
def export_docx():
    if not get_authenticated_user():
        return "Brak autoryzacji", 401
    
    content = request.form.get('content', '')
    title = request.form.get('title', 'Dokument')
    
    doc = Document()
    doc.add_heading(title, 0)
    
    for line in content.split('\n'):
        doc.add_paragraph(line)
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{title}.docx'
    )

@app.route('/export/pdf', methods=['POST'])
@csrf.exempt
def export_pdf():
    if not get_authenticated_user():
        return "Brak autoryzacji", 401
        
    content = request.form.get('content', '')
    title = request.form.get('title', 'Dokument')
    
    file_stream = io.BytesIO()
    pdf = canvas.Canvas(file_stream, pagesize=letter)
    
    try:
        sys_os = platform.system()
        if sys_os == "Windows":
            font_path = "C:\\Windows\\Fonts\\arial.ttf"
            font_path_bold = "C:\\Windows\\Fonts\\arialbd.ttf"
        elif sys_os == "Darwin":
            font_path = "/Library/Fonts/Arial.ttf"
            font_path_bold = "/Library/Fonts/Arial Bold.ttf"
        else:
            font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
            font_path_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

        pdfmetrics.registerFont(TTFont('PolishArial', font_path))
        pdfmetrics.registerFont(TTFont('PolishArial-Bold', font_path_bold))
        font_regular = 'PolishArial'
        font_bold = 'PolishArial-Bold'
    except:
        font_regular = 'Helvetica'
        font_bold = 'Helvetica-Bold'

    pdf.setTitle(title)
    
    pdf.setFont(font_bold, 16)
    pdf.drawString(50, 750, title)
    pdf.setStrokeColorRGB(0.2, 0.2, 0.2)
    pdf.line(50, 740, 550, 740)
    
    pdf.setFont(font_regular, 10)
    y = 710
    for line in content.split('\n'):
        clean_line = line.encode('utf-8', errors='ignore').decode('utf-8')
        wrapped_lines = simpleSplit(clean_line, font_regular, 10, 500) or [""]
        for wrapped_line in wrapped_lines:
            if y < 50:
                pdf.showPage()
                y = 750
                pdf.setFont(font_regular, 10)
            pdf.drawString(50, y, wrapped_line)
            y -= 15

        if not clean_line.strip():
            y -= 5

        if y < 50:
            pdf.showPage()
            y = 750
            pdf.setFont(font_regular, 10)
        
    pdf.save()
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'{title}.pdf'
    )


# ─── API Key management ────────────────────────────────────────────────────────
@app.route('/api/keys', methods=['GET'])
def list_api_keys():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = get_db_connection()
    rows = conn.execute(
        "SELECT id, name, allowed_origin, created_at, last_used_at FROM api_keys WHERE user_email = ? ORDER BY created_at DESC",
        (user_email,)
    ).fetchall()
    conn.close()
    return jsonify({"keys": [dict(r) for r in rows]})

@app.route('/api/keys', methods=['POST'])
@csrf.exempt
def create_api_key():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    name = str(data.get('name', '') or '').strip()[:80] or 'Klucz API'
    allowed_origin = str(data.get('allowed_origin', '') or '').strip()[:255]
    conn = get_db_connection()
    count = conn.execute("SELECT COUNT(*) FROM api_keys WHERE user_email = ?", (user_email,)).fetchone()[0]
    if count >= 10:
        conn.close()
        return jsonify({"error": "Maksymalna liczba kluczy API to 10"}), 400
    raw_key = 'stt_' + secrets.token_urlsafe(32)
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    cursor = conn.execute(
        "INSERT INTO api_keys (user_email, key_hash, name, allowed_origin) VALUES (?, ?, ?, ?)",
        (user_email, key_hash, name, allowed_origin)
    )
    key_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"id": key_id, "name": name, "allowed_origin": allowed_origin, "key": raw_key}), 201

@app.route('/api/keys/<int:key_id>', methods=['DELETE'])
@csrf.exempt
def delete_api_key(key_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = get_db_connection()
    result = conn.execute(
        "DELETE FROM api_keys WHERE id = ? AND user_email = ?", (key_id, user_email)
    )
    conn.commit()
    conn.close()
    return jsonify({"deleted": result.rowcount > 0})


@app.route('/api/compare-notes', methods=['POST'])
@limiter.limit("10 per hour; 3 per minute")
@csrf.exempt
def compare_notes_api():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    data = request.get_json(silent=True) or {}
    source_ids = data.get('source_ids') or []
    model_ids  = data.get('model_ids')  or []
    notes_mode = data.get('notes_mode', 'full')

    if not isinstance(model_ids, list) or not (1 <= len(model_ids) <= 4):
        return jsonify({"error": "Wybierz 1–4 modele"}), 400
    if notes_mode not in {'full', 'summary', 'overview', 'bullets', 'prompt'}:
        notes_mode = 'full'

    # load source texts
    raw_texts = []
    if source_ids:
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        for sid in source_ids[:5]:
            try:
                cursor.execute(
                    "SELECT filename, raw_text FROM history WHERE id = ? AND user_email = ?",
                    (int(sid), user_email)
                )
                row = cursor.fetchone()
                if row and row[1]:
                    raw_texts.append({"filename": row[0], "raw_text": row[1]})
            except (ValueError, TypeError):
                pass
        conn.close()

    if not raw_texts:
        return jsonify({"error": "Brak wybranych źródeł z tekstem"}), 400

    combined = build_project_context(raw_texts) if len(raw_texts) > 1 else raw_texts[0]['raw_text']

    # resolve model objects
    models_to_use = []
    for mid in model_ids[:4]:
        try:
            m = get_ai_model_by_id(int(mid), "notes")
            if m:
                models_to_use.append(m)
        except (ValueError, TypeError):
            pass
    if not models_to_use:
        return jsonify({"error": "Brak dostępnych modeli"}), 400

    import time
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def _generate(model):
        t0 = time.time()
        try:
            notes, _ = generate_audio_notes(
                combined, 'online',
                preferred_provider=model.get('provider'),
                model_used=None,
                notes_mode=notes_mode,
                notes_model_id=model.get('id')
            )
            return {
                "model_id": model["id"], "display_name": model["display_name"],
                "provider": model["provider"], "notes": notes,
                "timing_ms": int((time.time() - t0) * 1000), "error": None
            }
        except Exception as exc:
            return {
                "model_id": model["id"], "display_name": model["display_name"],
                "provider": model["provider"], "notes": None,
                "timing_ms": int((time.time() - t0) * 1000), "error": str(exc)
            }

    results = []
    with ThreadPoolExecutor(max_workers=min(4, len(models_to_use))) as pool:
        futures = {pool.submit(_generate, m): m for m in models_to_use}
        for f in as_completed(futures):
            results.append(f.result())

    order = {int(mid): i for i, mid in enumerate(model_ids)}
    results.sort(key=lambda r: order.get(r['model_id'], 999))

    # persist benchmark data
    src_chars = len(combined)
    tok_in_est = src_chars // 4
    conn2 = sqlite3.connect(DB_FILE)
    c2 = conn2.cursor()
    for r in results:
        tok_out_est = len(r.get('notes') or '') // 4
        c2.execute(
            "INSERT INTO compare_results "
            "(user_email, display_name, provider, notes_mode, source_chars, tokens_in_est, tokens_out_est, timing_ms, success) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (user_email, r['display_name'], r['provider'], notes_mode,
             src_chars, tok_in_est, tok_out_est, r['timing_ms'],
             0 if r['error'] else 1)
        )
    conn2.commit()
    conn2.close()

    sources_info = [{"id": s_id, "filename": t["filename"]}
                    for s_id, t in zip(source_ids, raw_texts)]

    return jsonify({"results": results, "sources": sources_info})


@app.route('/api/compare-reports', methods=['GET'])
@csrf.exempt
def compare_reports_api():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401

    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()

    cursor.execute("""
        SELECT display_name, provider,
               COUNT(*)                                    AS cnt,
               SUM(success)                               AS ok,
               ROUND(AVG(timing_ms))                      AS avg_ms,
               ROUND(MIN(timing_ms))                      AS min_ms,
               ROUND(MAX(timing_ms))                      AS max_ms,
               ROUND(AVG(tokens_in_est))                  AS avg_tok_in,
               ROUND(AVG(tokens_out_est))                 AS avg_tok_out,
               ROUND(AVG(CASE WHEN success=1 AND timing_ms > 0
                   THEN tokens_out_est * 1000.0 / timing_ms ELSE NULL END), 1) AS avg_tps
        FROM compare_results
        WHERE user_email = ?
        GROUP BY display_name, provider
        ORDER BY avg_ms ASC
    """, (user_email,))
    stats_rows = cursor.fetchall()
    stats_keys = ['display_name', 'provider', 'count', 'successes',
                  'avg_ms', 'min_ms', 'max_ms',
                  'avg_tokens_in', 'avg_tokens_out', 'avg_tps']
    stats = [dict(zip(stats_keys, r)) for r in stats_rows]

    cursor.execute("""
        SELECT id, display_name, provider, notes_mode, source_chars,
               tokens_in_est, tokens_out_est, timing_ms, success,
               datetime(created_at, 'localtime')
        FROM compare_results
        WHERE user_email = ?
        ORDER BY created_at DESC LIMIT 100
    """, (user_email,))
    runs_rows = cursor.fetchall()
    runs_keys = ['id', 'display_name', 'provider', 'notes_mode', 'source_chars',
                 'tokens_in', 'tokens_out', 'timing_ms', 'success', 'created_at']
    runs = [dict(zip(runs_keys, r)) for r in runs_rows]
    for r in runs:
        if r['timing_ms'] and r['success']:
            r['tps'] = round(r['tokens_out'] * 1000 / r['timing_ms'], 1) if r['timing_ms'] > 0 else 0
        else:
            r['tps'] = None

    conn.close()
    return jsonify({"stats": stats, "runs": runs})


@app.route('/api/compare-notes/recent-sources', methods=['GET'])
@csrf.exempt
def compare_recent_sources():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, filename, datetime(created_at,'localtime') FROM history "
        "WHERE user_email = ? AND raw_text != '' ORDER BY created_at DESC LIMIT 10",
        (user_email,)
    )
    rows = cursor.fetchall()
    conn.close()
    return jsonify([{"id": r[0], "filename": r[1], "created_at": r[2]} for r in rows])


@app.route('/api/projects', methods=['GET'])
@csrf.exempt
def list_projects():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT p.id, p.name, p.created_at,
               COUNT(h.id) AS source_count
        FROM projects p
        LEFT JOIN history h ON h.project_id = p.id
        WHERE p.user_email = ?
        GROUP BY p.id
        ORDER BY p.created_at DESC
        """,
        (user_email,)
    )
    rows = cursor.fetchall()
    conn.close()
    return jsonify([{"id": r[0], "name": r[1], "created_at": r[2], "source_count": r[3]} for r in rows])


@app.route('/api/projects', methods=['POST'])
@limiter.limit("60 per hour")
@csrf.exempt
def create_project():
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    name = str(data.get('name') or 'Nowy projekt').strip()[:200]
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO projects (user_email, name) VALUES (?, ?)", (user_email, name))
    project_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"id": project_id, "name": name}), 201


@app.route('/api/projects/<int:project_id>', methods=['PATCH'])
@csrf.exempt
def rename_project(project_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    name = str(data.get('name') or '').strip()[:200]
    if not name:
        return jsonify({"error": "Pusty tytuł"}), 400
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("UPDATE projects SET name = ? WHERE id = ? AND user_email = ?", (name, project_id, user_email))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
@csrf.exempt
def delete_project(project_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    # odepnij źródła (nie usuwa historii — tylko relację)
    cursor.execute("UPDATE history SET project_id = NULL WHERE project_id = ? AND user_email = ?", (project_id, user_email))
    cursor.execute("DELETE FROM projects WHERE id = ? AND user_email = ?", (project_id, user_email))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.route('/api/projects/<int:project_id>/sources', methods=['GET'])
@csrf.exempt
def list_project_sources(project_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM projects WHERE id = ? AND user_email = ?", (project_id, user_email))
    if not cursor.fetchone():
        conn.close()
        return jsonify({"error": "Nie znaleziono projektu"}), 404
    cursor.execute(
        "SELECT id, filename, datetime(created_at,'localtime') FROM history WHERE project_id = ? ORDER BY created_at",
        (project_id,)
    )
    rows = cursor.fetchall()
    conn.close()
    sources_count = len(rows)
    return jsonify({
        "project_id": project_id,
        "sources": [{"id": r[0], "filename": r[1], "created_at": r[2]} for r in rows],
        "context_warning": sources_count > 5,
    })


@app.route('/api/projects/<int:project_id>/sources', methods=['POST'])
@csrf.exempt
def add_project_source(project_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    data = request.get_json(silent=True) or {}
    history_id = data.get('history_id')
    if not history_id:
        return jsonify({"error": "Brak history_id"}), 400
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM projects WHERE id = ? AND user_email = ?", (project_id, user_email))
    if not cursor.fetchone():
        conn.close()
        return jsonify({"error": "Nie znaleziono projektu"}), 404
    cursor.execute("SELECT id FROM history WHERE id = ? AND user_email = ?", (history_id, user_email))
    if not cursor.fetchone():
        conn.close()
        return jsonify({"error": "Nie znaleziono nagrania"}), 404
    cursor.execute("UPDATE history SET project_id = ? WHERE id = ? AND user_email = ?", (project_id, history_id, user_email))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.route('/api/projects/<int:project_id>/sources/<int:history_id>', methods=['DELETE'])
@csrf.exempt
def remove_project_source(project_id, history_id):
    user_email = get_authenticated_user()
    if not user_email:
        return jsonify({"error": "Brak autoryzacji"}), 401
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM projects WHERE id = ? AND user_email = ?", (project_id, user_email))
    if not cursor.fetchone():
        conn.close()
        return jsonify({"error": "Nie znaleziono projektu"}), 404
    # odepnij źródło (nie tworzy nowego projektu — historia zostaje "osierocona")
    cursor.execute(
        "UPDATE history SET project_id = NULL WHERE id = ? AND project_id = ? AND user_email = ?",
        (history_id, project_id, user_email)
    )
    conn.commit()
    conn.close()
    return jsonify({"success": True})


if __name__ == '__main__':
    server_host = '0.0.0.0'
    server_port = get_server_port(host=server_host, default_port=8000)
    debug_mode = os.getenv('FLASK_DEBUG', '0').strip() == '1'
    app.run(debug=debug_mode, host=server_host, port=server_port)
